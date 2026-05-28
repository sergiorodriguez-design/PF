#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
conversor_papel.py v17-generico

Convierte una unidad online DOCX/PDF/TXT a DOCX papel editorial usando como base
un DOCX ya maquetado y, opcionalmente, un DOCX de interacciones.

Uso:
  python conversor_papel.py UNIDAD.docx EJEMPLO_MAQUETADO.docx SALIDA.docx
  python conversor_papel.py UNIDAD.docx EJEMPLO_MAQUETADO.docx INTERACCIONES.docx SALIDA.docx
  python conversor_papel.py UNIDAD.docx EJEMPLO_MAQUETADO.docx PLANTILLA.docx INTERACCIONES.docx SALIDA.docx

Cambios clave v17-generico:
  - Inserta gráficos reales del DOCX de referencia: imágenes, SmartArt, diagramas, dibujos.
  - Corrige el problema de namespace wp14: los SmartArt usaban wp14:anchorId y no se declaraba.
  - Copia TODO el paquete del ejemplo salvo document.xml, para que los rId y parts de SmartArt funcionen.
  - Extrae párrafos con regex correcta: no confunde <w:pPr> con <w:p>.
  - Usa los namespace reales del documento de referencia.
  - Conserva cursivas/negritas de los runs y muestra en papel la URL real de los hipervínculos.
  - Respeta la numeración visible del online en actividades y tareas.
  - Añade configuración genérica, versión alumno/docente y validación de salida.
"""

from __future__ import annotations

import re
import sys
import json
import zipfile
import subprocess
import tempfile
from pathlib import Path
from xml.sax.saxutils import escape as xml_escape


# =============================================================================
# Regex y constantes
# =============================================================================

RE_URL = re.compile(r"^https?://", re.IGNORECASE)
RE_SEC1 = re.compile(r"^(\d+)\.(?!\d)\s+(.+)")
RE_SEC2 = re.compile(r"^(\d+)\.(\d+)\.?\s+(.+)")
RE_SEC3 = re.compile(r"^(\d+)\.(\d+)\.(\d+)\.?\s+(.+)")
RE_INTER = re.compile(r"^Interacci[oó]n\s+(\d+)(?:\.?\s+(.+))?$", re.IGNORECASE)
RE_OPCION = re.compile(r"^([a-h])[\).]\s+(.+)")
RE_FORMULA = re.compile(
    r"(?:"
    r"\d+\s*[×x\*/÷]\s*\d+"
    r"|[A-Za-zÁÉÍÓÚÜÑáéíóúüñ]+\s*=\s*.+[/÷×]"
    r"|\b\w+\s*/\s*\w+\s*[×x]\s*100"
    r")"
)

RE_SOLUCION_MARKER = re.compile(
    r"^(?:POSIBLE\s+SOLUCI[OÓ]N|SOLUCI[OÓ]N(?:\s*\([^)]*posible[^)]*\))?|Soluci[oó]n|Feedback|Retroalimentaci[oó]n)\b",
    re.I,
)

def _es_marcador_solucion(texto: str) -> bool:
    return bool(RE_SOLUCION_MARKER.match((texto or "").strip()))

def _bloque_admite_solucion(blk: dict | None) -> bool:
    """Regla editorial: la solución solo se conserva en Aplicación práctica."""
    if not blk:
        return False
    etiqueta = str(blk.get("etiqueta", "")).strip().lower()
    tipo = str(blk.get("tipo", "")).strip().lower()
    return tipo == "aplicacion_practica" or etiqueta.startswith("aplicación práctica") or etiqueta.startswith("aplicacion practica")

RE_PUA = re.compile(r"[\ue000-\uf8ff\U000f0000-\U000fffff]+")
RE_PUA_ONLY = re.compile(r"^[\ue000-\uf8ff\U000f0000-\U000fffff\s]+$")

BLOQUES_ESP = {
    "Nota", "Ejemplo", "Sabías que...", "Sabías que…", "Consejo",
    "Definición", "Hilo conductor", "Para saber más", "Vídeo", "Video",
    "Importante", "Recuerda", "Actividad complementaria", "Actividad colaborativa",
    "Actividad de evaluación", "Aplicación práctica", "Caso práctico", "Ejercicio", "Tarea",
}

PREFIJOS_ELIM = (
    "Para realizar las Actividades colaborativas",
    "Para realizar las Tareas de evaluación",
    "Las instrucciones para realizar la tarea",
    "Podrás compartir ",
    "Podrás debatir",
    "poder debatir y aportar",
    "Podrás identificar las",
    "Es el momento de realizar la siguiente",
    "No obstante puedes seguir estudiando",
    "La duración aproximada de la misma",
    "Mapa conceptual o esquema de contenidos",
    "Para realizar las",
    "valoración será tenida en cuenta",
    "encontrarás la información necesaria",
    "otro momento que te sea más favorable",
    "Cambio de pantalla",
    "Pulsa en ",
    "Pulsa para ",
    "Haz clic ",
    "Haz clic para ",
    "Haz clic en ",
    "Pincha ",
    "Instrucción:",
    "Instrucción: ",
    "Desplegables:",
    "Duración:",
    "Duracion:",
    "Criterios de evaluación:",
    "Criterios de evaluacion:",
    "Lanzador:",
)

RESIDUOS_RESUMEN = {
    "Conservación o eliminación",
    "Almacenamiento y custodia",
    "Uso y tramitación",
    "Clasificación o registro",
    "Creación o recepción",
    "Clasificación y registro",
}

VINETA_SIM = {1: "●", 2: "○", 3: "▪", 4: "–"}
VINETA_EST = {
    1: "Vietanvl11d",
    2: "Vietanvl21d",
    3: "Vietanvl31d",
    4: "Vietanvl41d",
}

DEFAULT_CONFIG = {
    "document_type": "unidad_aprendizaje",
    "version": "alumno",
    "unit_labels": ["Unidad de aprendizaje", "Unidad didáctica", "Tema", "Módulo", "UA"],
    "remove_labels": [
        "Cambio de pantalla", "Criterios de evaluación", "Criterios de evaluacion",
        "Duración", "Duracion", "Instrucción", "Instruccion", "Desplegables"
    ],
    "box_labels": sorted(BLOQUES_ESP),
    "activity_label_map": {
        "Actividad colaborativa": "Actividad complementaria",
        "Actividad de evaluación": "Actividad",
    },
    "interaction_mode": "expand",
    "include_solutions": False,
    "include_feedback": False,
    "objectives_intro": "Los objetivos específicos de esta Unidad de Aprendizaje son:",
    "images": {"embed": False, "keep_url": True, "keep_caption": True, "keep_description": True},
    "bullets": {"use_word_lists": False, "level_1": "●", "level_2": "○"},
    "highlight": {"simple_interactions": True, "skip_complex_interactions": True, "color": "yellow", "max_words_per_item": 80},
    "validation": {"fail_on_errors": False},
}

RE_REMOVE_LABEL_ONLY = re.compile(
    r"^(?:Cambio de pantalla|Criterios de evaluaci[oó]n|Duraci[oó]n|Instrucci[oó]n|Desplegables|Opciones|Enunciado)\s*:?\s*$",
    re.I,
)
RE_REMOVE_LABEL_PREFIX = re.compile(
    r"^(?:Enunciado|Instrucci[oó]n|Opciones)\s*:\s*(.+)$",
    re.I,
)

RE_GENERIC_UNIT = re.compile(
    r"^(Unidad\s+de\s+aprendizaje|Unidad\s+did[aá]ctica|Unidad|Tema|M[oó]dulo|UA)\s+(\d+)(?:(?:\s*[.\-–:]\s*|\s+)(.+))?$",
    re.I,
)

RE_ACTIVITY_LABEL = re.compile(
    r"^(Actividad(?:\s+(?:complementaria|colaborativa|de evaluaci[oó]n|de aprendizaje))?|Aplicaci[oó]n pr[aá]ctica|Caso pr[aá]ctico|Ejercicio|Tarea)(?:\s+\d+)?$",
    re.I,
)

RE_FORBIDDEN_COMMON = re.compile(r"\b(Cambio de pantalla|Desplegables:|Instrucci[oó]n:|Interacci[oó]n\s+\d+)\b", re.I)
RE_FORBIDDEN_STUDENT = re.compile(r"\b(Duraci[oó]n:|Soluci[oó]n:|Feedback:|Retroalimentaci[oó]n:|Criterios de evaluaci[oó]n:)\b", re.I)

RUNTIME_CONFIG = dict(DEFAULT_CONFIG)


def _deep_update(base: dict, extra: dict) -> dict:
    out = dict(base)
    for k, v in (extra or {}).items():
        if isinstance(v, dict) and isinstance(out.get(k), dict):
            out[k] = _deep_update(out[k], v)
        else:
            out[k] = v
    return out


def cargar_config(path: Path | None = None, version: str | None = None) -> dict:
    cfg = dict(DEFAULT_CONFIG)
    raw = {}
    if path:
        raw = json.loads(path.read_text(encoding="utf-8"))
        cfg = _deep_update(cfg, raw)
    if version:
        cfg["version"] = version
    # Por defecto, docente conserva soluciones/feedback; si el JSON los define explícitamente, se respeta.
    if cfg.get("version") == "docente":
        if "include_solutions" not in raw:
            cfg["include_solutions"] = True
        if "include_feedback" not in raw:
            cfg["include_feedback"] = True
    else:
        if "include_solutions" not in raw:
            cfg["include_solutions"] = False
        if "include_feedback" not in raw:
            cfg["include_feedback"] = False
    return cfg


def set_runtime_config(cfg: dict) -> None:
    global RUNTIME_CONFIG, BLOQUES_ESP
    RUNTIME_CONFIG = _deep_update(DEFAULT_CONFIG, cfg or {})
    BLOQUES_ESP = set(RUNTIME_CONFIG.get("box_labels") or BLOQUES_ESP)


def normalizar_etiqueta_actividad(label: str, cfg: dict | None = None) -> str:
    cfg = cfg or RUNTIME_CONFIG
    label = (label or "").strip()
    return (cfg.get("activity_label_map") or {}).get(label, label)


def limpiar_metadato_generico(texto: str) -> str | None:
    """Devuelve None si la línea debe eliminarse; devuelve texto limpio si solo se quita la etiqueta."""
    t = _norm_line(texto) if "_norm_line" in globals() else (texto or "").strip()
    if not t:
        return ""
    if RE_REMOVE_LABEL_ONLY.match(t):
        return None
    m = RE_REMOVE_LABEL_PREFIX.match(t)
    if m:
        return m.group(1).strip()
    if RUNTIME_CONFIG.get("version") == "alumno" and re.match(r"^(Soluci[oó]n|Feedback|Retroalimentaci[oó]n|Duraci[oó]n|Criterios de evaluaci[oó]n)\s*:", t, re.I):
        return None
    return t


# =============================================================================
# Helpers generales
# =============================================================================

def esc(texto: str) -> str:
    return xml_escape(str(texto), {'"': "&quot;"})


def limpiar_pua(texto: str) -> str:
    return RE_PUA.sub("", texto or "").strip()


def es_solo_pua(texto: str) -> bool:
    return bool(texto) and bool(RE_PUA_ONLY.match(texto))


# Verbos/frases genéricas de navegación que van tras "Avanza para "
_RE_AVANZA_NAV = re.compile(
    r"^(continuar|ver|pasar|seguir|empezar|comenzar|acceder|avanzar|terminar|"
    r"el siguiente|la siguiente|los siguientes|las siguientes|ir a|ir al)\b",
    re.I,
)


def debe_elim(texto: str) -> bool:
    if not texto:
        return True
    limpio = limpiar_metadato_generico(texto)
    if limpio is None:
        return True
    if texto in RESIDUOS_RESUMEN:
        return True
    if texto.startswith("Avanza para "):
        resto = texto[len("Avanza para "):]
        return bool(_RE_AVANZA_NAV.match(resto))
    return any(texto.startswith(p) for p in PREFIJOS_ELIM)


def limpiar_titulo(texto: str) -> str:
    texto = re.sub(r"\s*\(CE[^)]*\)", "", texto, flags=re.I)
    texto = re.sub(r"\s+CE\s+[a-z](?:\s*[,y]\s*CE\s+[a-z])*\s*$", "", texto, flags=re.I)
    texto = re.sub(r"\s*…+$", "", texto)
    texto = re.sub(r"\s*\.{2,}$", "", texto)
    return texto.rstrip("…. ").strip()




def _normalizar_heading_texto(texto: str) -> str:
    """Limpia marcas frecuentes de CE y puntos sobrantes en títulos."""
    return limpiar_titulo(texto or "")


def _match_titulo_unidad(texto: str):
    """Acepta etiquetas de unidad parametrizables: Unidad de aprendizaje, Unidad didáctica, Tema, Módulo, UA."""
    texto = _norm_line(texto)
    m = RE_GENERIC_UNIT.match(texto)
    if not m:
        return None
    label = re.sub(r"\s+", " ", m.group(1)).strip()
    # Normaliza capitalización sin depender del tema ni del número.
    label_norm = {
        "unidad de aprendizaje": "Unidad de aprendizaje",
        "unidad didáctica": "Unidad didáctica",
        "unidad didactica": "Unidad didáctica",
        "tema": "Tema",
        "módulo": "Módulo",
        "modulo": "Módulo",
        "ua": "UA",
    }.get(label.lower(), label[:1].upper() + label[1:])
    titulo = (m.group(3) or "").strip()
    return label_norm, int(m.group(2)), limpiar_titulo(titulo)

def _es_cabecera_objetivos(texto: str) -> bool:
    t = _norm_line(texto).lower().strip(':')
    return t in {
        "objetivos", "objetivos específicos", "objetivos especificos",
        "objetivos generales", "objetivo general",
        "resultados de aprendizaje", "objetivos de aprendizaje",
    } or t.startswith("los objetivos específicos") or t.startswith("los objetivos especificos")


def _tipo_objetivos_intro(texto: str) -> str:
    """Devuelve el texto de introducción correcto según el encabezado de objetivos."""
    t = _norm_line(texto).lower().strip(':')
    if "generales" in t or t == "objetivo general":
        return "Los objetivos generales de esta Unidad de Aprendizaje son:"
    return "Los objetivos específicos de esta Unidad de Aprendizaje son:"


def _es_cabecera_no_contenido(texto: str) -> bool:
    t = _norm_line(texto).lower().strip(':')
    return t in {
        "resultado de aprendizaje", "resultados de aprendizaje",
        "criterios de evaluación", "criterios de evaluacion",
        "ce", "ra",
    }

def infinitivo_a_imperativo(texto: str) -> str:
    pares = [
        ("buscar ", "Busca "), ("identificar ", "Identifica "),
        ("analizar ", "Analiza "), ("elaborar ", "Elabora "),
        ("diseñar ", "Diseña "), ("crear ", "Crea "),
        ("realizar ", "Realiza "), ("comparar ", "Compara "),
        ("describir ", "Describe "), ("explicar ", "Explica "),
        ("completar ", "Completa "), ("redactar ", "Redacta "),
        ("investigar ", "Investiga "), ("seleccionar ", "Selecciona "),
        ("clasificar ", "Clasifica "), ("calcular ", "Calcula "),
        ("consultar ", "Consulta "), ("revisar ", "Revisa "),
        ("explorar ", "Explora "), ("reflexionar ", "Reflexiona "),
        ("compartir ", "Comparte "), ("debatir ", "Debate "),
        ("aplicar ", "Aplica "), ("desarrollar ", "Desarrolla "),
        ("resolver ", "Resuelve "), ("relacionar ", "Relaciona "),
        ("evaluar ", "Evalúa "), ("definir ", "Define "),
        ("justificar ", "Justifica "), ("plantear ", "Plantea "),
        ("proponer ", "Propón "), ("observar ", "Observa "),
    ]
    low = texto.lower()
    for inf, imp in pares:
        if low.startswith(inf):
            return imp + texto[len(inf):]
    return texto[:1].upper() + texto[1:] if texto else texto


def _futuro_a_imperativo(verbo: str) -> str:
    """Convierte un verbo en futuro 2ª persona (tú) a imperativo.
    explorarás → Explora, aprenderás → Aprende, escribirás → Escribe"""
    v = verbo.lower().rstrip(".,;:")
    if v.endswith("arás"):
        imp = v[:-4] + "a"
    elif v.endswith("erás"):
        imp = v[:-4] + "e"
    elif v.endswith("irás"):
        imp = v[:-4] + "e"
    else:
        return verbo[0].upper() + verbo[1:] if verbo else verbo
    return imp[0].upper() + imp[1:]


def _es_url_imagen(url: str) -> bool:
    dominios = (
        "shutterstock.com", "gettyimages.", "istockphoto.com", "unsplash.com",
        "pexels.com", "freepik.com", "pixabay.com", "depositphotos.com",
        "adobe.com/stock", "stock.adobe.com",
    )
    u = url.lower()
    if any(d in u for d in dominios):
        return True
    return bool(re.search(r"\.(jpg|jpeg|png|gif|webp|svg|bmp|tiff?)(\?|$)", u))


def _es_zip(path: Path) -> bool:
    try:
        with open(path, "rb") as f:
            return f.read(4) == b"PK\x03\x04"
    except Exception:
        return False


def _style_name(p) -> str:
    try:
        return p.style.name
    except Exception:
        return ""


def _is_bold(p) -> bool:
    return any(r.bold for r in p.runs if r.text.strip())


def _norm_line(texto: str) -> str:
    texto = limpiar_pua(texto)
    texto = texto.replace("\x00", "")
    texto = re.sub(r"\s+", " ", texto)
    return texto.strip()


def _limpiar_vineta_literal(texto: str) -> str:
    texto = texto.strip()
    texto = re.sub(r"^[●○▪·•*\-–]\s*", "", texto)
    texto = re.sub(r"^[●○▪·•*]\t", "", texto)
    return texto.strip()


def _abre_modo_lista(linea: str) -> bool:
    l = linea.strip().lower()
    if not l.endswith(":"):
        return False

    claves = (
        "permite", "mediante", "divide", "dividirse", "aspectos", "como",
        "son", "serían", "radica en", "se basa en", "incluye", "destacan",
        "siguientes", "pudiendo ser",
    )
    return any(c in l for c in claves)


def _parece_item_lista_en_bloque(linea: str, modo_lista: bool) -> bool:
    linea = linea.strip()

    if not linea:
        return False

    if re.match(r"^[●○▪\-–]\s+", linea):
        return True

    if not modo_lista:
        return False

    if len(linea) > 110:
        return False

    if linea.startswith((
        "Algunas ", "Alguna ", "En ", "Por ", "Para ", "Cuando ",
        "Una ", "Un ", "El ", "La ", "Los ", "Las ", "Este ", "Esta ",
        "Estos ", "Estas ", "Según ", "Asimismo ", "Además ",
    )):
        return False

    return True



# =============================================================================
# Texto enriquecido e hipervínculos
# =============================================================================

_HYPERLINKS_OUT: dict[str, str] = {}
_HYPERLINK_SEQ = 1


def _reset_hyperlinks_out() -> None:
    global _HYPERLINKS_OUT, _HYPERLINK_SEQ
    _HYPERLINKS_OUT = {}
    _HYPERLINK_SEQ = 1


def _rid_hyperlink(target: str) -> str:
    """Devuelve/crea un rId estable para un hipervínculo externo de salida."""
    global _HYPERLINK_SEQ
    target = str(target or "").strip()
    if not target:
        return ""
    if target not in _HYPERLINKS_OUT:
        _HYPERLINKS_OUT[target] = f"rIdHyper{_HYPERLINK_SEQ}"
        _HYPERLINK_SEQ += 1
    return _HYPERLINKS_OUT[target]


def rich_text(obj) -> str:
    if isinstance(obj, dict):
        return str(obj.get("texto", ""))
    return "" if obj is None else str(obj)


def rich_runs(obj):
    if isinstance(obj, dict):
        return obj.get("runs") or [{"text": rich_text(obj)}]
    return [{"text": rich_text(obj)}]


def rich_obj(texto: str, runs=None) -> dict:
    return {"texto": texto or "", "runs": runs or [{"text": texto or ""}]}


def _trim_runs(runs: list[dict]) -> list[dict]:
    runs = [dict(r) for r in (runs or []) if str(r.get("text", ""))]
    if not runs:
        return []
    runs[0]["text"] = str(runs[0].get("text", "")).lstrip()
    runs[-1]["text"] = str(runs[-1].get("text", "")).rstrip()
    return [r for r in runs if str(r.get("text", ""))]


def _materializar_hipervinculos_para_papel(runs: list[dict]) -> list[dict]:
    """
    En papel no debe quedar el texto ancla del online (p. ej. el título de un vídeo),
    sino la URL real. Conserva el hipervínculo en el DOCX generado y muestra la URL
    como texto visible.
    """
    out: list[dict] = []
    last_link = None

    for r in runs or []:
        rr = dict(r)
        link = str(rr.get("link", "")).strip()

        if link and RE_URL.match(link):
            # Si varios runs consecutivos forman el mismo hipervínculo, se sustituyen
            # por una sola URL visible para evitar duplicados.
            if link == last_link:
                continue
            rr["text"] = link
            rr["link"] = link
            # El estilo de negrita/cursiva del texto ancla no debe contaminar la URL.
            rr.pop("bold", None)
            rr.pop("italic", None)
            out.append(rr)
            last_link = link
        else:
            out.append(rr)
            last_link = None

    return _trim_runs(out)


def _runs_to_xml(runs: list[dict], force_bold: bool = False) -> str:
    partes: list[str] = []
    for r in _trim_runs(runs):
        txt = str(r.get("text", ""))
        if not txt:
            continue
        te = esc(txt)
        sp = ' xml:space="preserve"' if te and te != te.strip() else ""
        props = []
        if r.get("link"):
            props.append('<w:rStyle w:val="Hyperlink"/>')
        if force_bold or r.get("bold"):
            props.append("<w:b/><w:bCs/>")
        if r.get("italic"):
            props.append("<w:i/><w:iCs/>")
        color = str(r.get("color", "")).strip()
        if color and re.match(r"^[0-9A-Fa-f]{6}$", color):
            props.append(f'<w:color w:val="{color.upper()}"/>')
        rpr = f"<w:rPr>{''.join(props)}</w:rPr>" if props else ""
        run_xml = f"<w:r>{rpr}<w:t{sp}>{te}</w:t></w:r>"
        if r.get("link"):
            rid = _rid_hyperlink(str(r.get("link")))
            if rid:
                run_xml = f'<w:hyperlink r:id="{rid}" w:history="1">{run_xml}</w:hyperlink>'
        partes.append(run_xml)
    return "".join(partes)


def p_rich(obj, estilo: str, negrita: bool = False) -> str:
    texto = rich_text(obj)
    runs = rich_runs(obj)
    if not texto and not runs:
        return f'    <w:p><w:pPr><w:pStyle w:val="{estilo}"/></w:pPr></w:p>'
    return (
        f'    <w:p><w:pPr><w:pStyle w:val="{estilo}"/></w:pPr>'
        f"{_runs_to_xml(runs, force_bold=negrita)}</w:p>"
    )


def _map_runs_text(runs: list[dict], func) -> dict:
    texto = func("".join(str(r.get("text", "")) for r in runs))
    return rich_obj(texto, [{"text": texto}])


def _strip_literal_from_rich(obj, literal: str) -> dict:
    texto = rich_text(obj)
    runs = rich_runs(obj)
    if not texto.startswith(literal):
        return rich_obj(texto, runs)
    rest = texto[len(literal):].lstrip()
    return rich_obj(rest, [{"text": rest}])


def _titulo_con_colon(titulo):
    """Strip trailing punctuation from a title and append ':'."""
    if isinstance(titulo, dict):
        runs = [dict(r) for r in rich_runs(titulo)]
        if runs:
            runs[-1]["text"] = runs[-1].get("text", "").rstrip(".: ") + ":"
        texto = rich_text(titulo).rstrip(".: ") + ":"
        return rich_obj(texto, runs)
    return str(titulo).rstrip(".: ") + ":"


def _prefix_rich(prefix: str, obj) -> dict:
    runs = [{"text": prefix}] + rich_runs(obj)
    result = rich_obj(prefix + rich_text(obj), runs)
    if isinstance(obj, dict) and obj.get("links"):
        result["links"] = obj["links"]
    return result


def _limpiar_vineta_rich(obj) -> dict:
    texto = rich_text(obj)
    limpio = _limpiar_vineta_literal(texto)
    if limpio == texto:
        return rich_obj(texto, rich_runs(obj))
    # Strip bullet prefix from the first run's text, preserving bold/italic of remaining runs.
    runs = [dict(r) for r in rich_runs(obj)]
    for i, r in enumerate(runs):
        t = r.get("text", "")
        t_clean = _limpiar_vineta_literal(t)
        if t_clean != t:
            runs[i]["text"] = t_clean
            break
    runs = [r for r in runs if r.get("text", "")]
    return rich_obj(limpio, runs or [{"text": limpio}])


def _patch_document_rels_hyperlinks(rels_xml: bytes | str) -> bytes:
    xml = rels_xml.decode("utf-8") if isinstance(rels_xml, bytes) else str(rels_xml)
    if not _HYPERLINKS_OUT:
        return xml.encode("utf-8")
    inserts = []
    for target, rid in _HYPERLINKS_OUT.items():
        if f'Id="{rid}"' in xml:
            continue
        inserts.append(
            f'<Relationship Id="{rid}" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink" '
            f'Target="{esc(target)}" TargetMode="External"/>'
        )
    if inserts:
        xml = xml.replace("</Relationships>", "".join(inserts) + "</Relationships>")
    return xml.encode("utf-8")

# =============================================================================
# Constructores XML
# =============================================================================

def p(texto: str, estilo: str, negrita: bool = False) -> str:
    if isinstance(texto, dict):
        return p_rich(texto, estilo, negrita)

    texto = "" if texto is None else str(texto)
    te = esc(texto)
    sp = ' xml:space="preserve"' if te and te != te.strip() else ""

    if not te:
        return f'    <w:p><w:pPr><w:pStyle w:val="{estilo}"/></w:pPr></w:p>'

    rpr = "<w:rPr><w:b/><w:bCs/></w:rPr>" if negrita else ""

    return (
        f'    <w:p><w:pPr><w:pStyle w:val="{estilo}"/></w:pPr>'
        f"<w:r>{rpr}<w:t{sp}>{te}</w:t></w:r></w:p>"
    )


def p_vineta(texto: str, nivel: int = 1) -> str:
    estilo = VINETA_EST.get(nivel, VINETA_EST[1])
    simbolo = esc(VINETA_SIM.get(nivel, VINETA_SIM[1]))

    if isinstance(texto, dict):
        cleaned = _limpiar_vineta_rich(texto)
        runs = rich_runs(cleaned)
        return (
            f'    <w:p><w:pPr><w:pStyle w:val="{estilo}"/></w:pPr>'
            f"<w:r><w:t>{simbolo}</w:t></w:r>"
            f"<w:r><w:tab/></w:r>{_runs_to_xml(runs)}</w:p>"
        )

    te = esc(_limpiar_vineta_literal(texto))

    return (
        f'    <w:p><w:pPr><w:pStyle w:val="{estilo}"/></w:pPr>'
        f"<w:r><w:t>{simbolo}</w:t></w:r>"
        f"<w:r><w:tab/><w:t>{te}</w:t></w:r></w:p>"
    )


def p_vineta_ejemplo(texto: str) -> str:
    if isinstance(texto, dict):
        cleaned = _limpiar_vineta_rich(texto)
        runs = rich_runs(cleaned)
        return (
            f'    <w:p><w:pPr><w:pStyle w:val="Ejemplos-Vietanvl1"/></w:pPr>'
            f"<w:r><w:t>●</w:t></w:r>"
            f"<w:r><w:tab/></w:r>{_runs_to_xml(runs)}</w:p>"
        )

    te = esc(_limpiar_vineta_literal(texto))

    return (
        f'    <w:p><w:pPr><w:pStyle w:val="Ejemplos-Vietanvl1"/></w:pPr>'
        f"<w:r><w:t>●</w:t></w:r>"
        f"<w:r><w:tab/><w:t>{te}</w:t></w:r></w:p>"
    )


def p_vineta_recuerda(texto) -> str:
    if isinstance(texto, dict):
        cleaned = _limpiar_vineta_rich(texto)
        runs = rich_runs(cleaned)
        return (
            f'    <w:p><w:pPr><w:pStyle w:val="Recuerda-Vietanvl1"/></w:pPr>'
            f"<w:r><w:t>●</w:t></w:r>"
            f"<w:r><w:tab/></w:r>{_runs_to_xml(runs)}</w:p>"
        )
    te = esc(_limpiar_vineta_literal(texto))
    return (
        f'    <w:p><w:pPr><w:pStyle w:val="Recuerda-Vietanvl1"/></w:pPr>'
        f"<w:r><w:t>●</w:t></w:r>"
        f"<w:r><w:tab/><w:t>{te}</w:t></w:r></w:p>"
    )


def p_vineta_bold(texto, nivel: int = 1) -> str:
    estilo = VINETA_EST.get(nivel, VINETA_EST[1])
    simbolo = esc(VINETA_SIM.get(nivel, VINETA_SIM[1]))

    if isinstance(texto, dict):
        cleaned = _limpiar_vineta_rich(texto)
        runs = []
        for r in rich_runs(cleaned):
            nr = dict(r)
            nr["bold"] = True
            runs.append(nr)
        return (
            f'    <w:p><w:pPr><w:pStyle w:val="{estilo}"/></w:pPr>'
            f"<w:r><w:t>{simbolo}</w:t></w:r>"
            f"<w:r><w:tab/></w:r>{_runs_to_xml(runs)}</w:p>"
        )

    te = esc(_limpiar_vineta_literal(texto))
    return (
        f'    <w:p><w:pPr><w:pStyle w:val="{estilo}"/></w:pPr>'
        f"<w:r><w:t>{simbolo}</w:t></w:r>"
        f"<w:r><w:tab/></w:r>"
        f"<w:r><w:rPr><w:b/><w:bCs/></w:rPr><w:t>{te}</w:t></w:r></w:p>"
    )


def p_desp(titulo, desc, nivel: int = 1) -> str:
    estilo = VINETA_EST.get(nivel, VINETA_EST[1])
    simbolo = esc(VINETA_SIM.get(nivel, VINETA_SIM[1]))
    titulo_plain = ""
    if isinstance(titulo, dict):
        titulo_plain = rich_text(titulo).rstrip(".: ")
        t_runs = list(rich_runs(titulo))
        if t_runs:
            last = dict(t_runs[-1])
            last["text"] = last.get("text", "").rstrip(".: ")
            t_runs = list(t_runs[:-1]) + [last]
        titulo_block = _runs_to_xml(t_runs) + f'<w:r><w:rPr><w:b/><w:bCs/></w:rPr><w:t xml:space="preserve">: </w:t></w:r>'
    else:
        titulo_str = str(titulo).rstrip(".: ")
        titulo_plain = titulo_str
        titulo_block = f'<w:r><w:rPr><w:b/><w:bCs/></w:rPr><w:t xml:space="preserve">{esc(titulo_str)}: </w:t></w:r>'
    header = (
        f'    <w:p><w:pPr><w:pStyle w:val="{estilo}"/></w:pPr>'
        f"<w:r><w:t>{simbolo}</w:t></w:r>"
        f"<w:r><w:tab/></w:r>"
        f"{titulo_block}"
    )
    if isinstance(desc, dict):
        desc_text = rich_text(desc).strip()
        runs = [dict(r) for r in rich_runs(desc)]
        if titulo_plain and desc_text.lower().startswith(titulo_plain.lower()):
            m = re.match(re.escape(titulo_plain) + r"[\.\s:]*", desc_text, re.I)
            if m:
                to_remove = m.end()
                removed = 0
                for i, r in enumerate(runs):
                    if removed >= to_remove:
                        break
                    t = r.get("text", "")
                    if len(t) <= (to_remove - removed):
                        removed += len(t)
                        r["text"] = ""
                    else:
                        r["text"] = t[(to_remove - removed):]
                        removed = to_remove
        for r in runs:
            if r.get("text"):
                r["text"] = r["text"][:1].lower() + r["text"][1:]
                break
        return header + _runs_to_xml(runs) + "</w:p>"
    desc_str = re.sub(r"::+", ":", str(desc).strip())
    if titulo_plain and desc_str.lower().startswith(titulo_plain.lower()):
        m = re.match(re.escape(titulo_plain) + r"[\.\s:]*", desc_str, re.I)
        if m:
            desc_str = desc_str[m.end():]
    if desc_str:
        desc_str = desc_str[:1].lower() + desc_str[1:]
    return header + f"<w:r><w:t>{esc(desc_str)}</w:t></w:r></w:p>"


def p_formula(texto: str) -> str:
    return p(texto, "Formula")


def p_opcion_test(letra: str, texto) -> str:
    if isinstance(texto, dict):
        cleaned = _limpiar_vineta_rich(texto)
        runs = rich_runs(cleaned)
        return (
            '    <w:p><w:pPr><w:pStyle w:val="EjerciciosPregunta"/></w:pPr>'
            f'<w:r><w:t xml:space="preserve">{esc(letra)}) </w:t></w:r>'
            f"{_runs_to_xml(runs)}</w:p>"
        )
    texto = re.sub(r"^[a-h][\).]\s*", "", str(texto).strip())
    return (
        '    <w:p><w:pPr><w:pStyle w:val="EjerciciosPregunta"/></w:pPr>'
        f"<w:r><w:t>{esc(letra)}) {esc(texto)}</w:t></w:r>"
        "</w:p>"
    )


def p_url_imagen(url: str) -> str:
    return p(rich_obj(url, [{"text": url, "link": url}]), "Cuerpoparrafo")


def p_url_recurso(url: str) -> str:
    return p(rich_obj(url, [{"text": url, "link": url}]), "Cuerpoparrafo")


def _links_from_rich(obj) -> list[str]:
    if not isinstance(obj, dict):
        return []
    links = []
    for link in obj.get("links", []) or []:
        link = str(link).strip()
        if link and RE_URL.match(link) and link not in links:
            links.append(link)
    return links


def _append_links_xml(out: list[str], obj) -> None:
    # En papel los enlaces deben quedar visibles al final del párrafo, en línea
    # aparte. No se sustituyen dentro del texto corrido.
    for link in _links_from_rich(obj):
        out.append(p_url_recurso(link))

def _texto_plano_runs(runs: list[dict]) -> str:
    return "".join(str(r.get("text", "")) for r in runs or [])


def _replace_text_preserve_first_style(obj, nuevo: str) -> dict:
    """Sustituye el texto visible conservando, si puede, el estilo del primer run."""
    links = _links_from_rich(obj)
    first = {}
    if isinstance(obj, dict) and obj.get("runs"):
        first = {k: v for k, v in dict(obj["runs"][0]).items() if k in {"bold", "italic", "color", "link"}}
    first["text"] = nuevo
    out = rich_obj(nuevo, [first])
    if links:
        out["links"] = links
    return out


def _normalizar_enunciado_complementaria(obj) -> dict:
    """Convierte redacciones online a enunciado papel.
    'En esta actividad deberás pensar...' → 'Piensa...'
    'En esta actividad explorarás la guía...' → 'Explora la guía...'
    """
    texto = rich_text(obj).strip()
    if not texto:
        return obj

    nuevo = texto

    # 1. Patrones con "deberás" + infinitivo
    reglas_debers = [
        (r"^En esta actividad deberás pensar en\s+", "Piensa en "),
        (r"^En esta actividad deberás\s+", ""),
        (r"^Deberás realizar\s+", "Realiza "),
        (r"^Deberás\s+", ""),
    ]
    for pat, rep in reglas_debers:
        nuevo2 = re.sub(pat, rep, nuevo, flags=re.I)
        if nuevo2 != nuevo:
            nuevo = infinitivo_a_imperativo(nuevo2.strip())
            break

    # 2. "En esta actividad VERB_FUTURO(rás) ..." → imperativo
    if nuevo == texto:
        m_fut = re.match(r"^En esta actividad[,]?\s+(\w+rás)[,]?\s*(.*)", nuevo, re.I | re.S)
        if m_fut:
            verbo = m_fut.group(1)
            resto = m_fut.group(2).strip()
            imp = _futuro_a_imperativo(verbo)
            nuevo = (imp + " " + resto).strip() if resto else imp

    # 3. "En esta actividad VERB_INFINITIVO ..." → imperativo
    if nuevo == texto:
        m_inf = re.match(r"^En esta actividad[,]?\s+(\w+[aei]r)\s+(.*)", nuevo, re.I | re.S)
        if m_inf:
            verbo = m_inf.group(1)
            resto = m_inf.group(2).strip()
            imp_txt = infinitivo_a_imperativo(verbo + " " + resto)
            if imp_txt != verbo + " " + resto:
                nuevo = imp_txt

    # 4. Frases de colaboración a eliminar
    for pat in (r"^Podrás compartir.*$", r"^Podrás debatir.*$", r"^Podrás aportar.*$"):
        if re.match(pat, nuevo, re.I):
            nuevo = ""
            break

    nuevo = re.sub(r"\s+", " ", nuevo).strip()
    if not nuevo:
        return rich_obj("", [])
    if nuevo == texto:
        return obj
    return _replace_text_preserve_first_style(obj, nuevo)


def _solucion_feedback_a_lineas(sol, fb) -> list[dict]:
    lineas = []
    if isinstance(sol, dict):
        sol_txt = rich_text(sol).strip()
        if sol_txt:
            lineas.append(rich_obj(
                f"Solución: {sol_txt}",
                [{"text": "Solución: ", "bold": True}] + list(rich_runs(sol))
            ))
    else:
        sol_str = (sol or "").strip()
        if sol_str:
            lineas.append(rich_obj(f"Solución: {sol_str}", [{"text": "Solución: ", "bold": True}, {"text": sol_str}]))
    if isinstance(fb, dict):
        fb_txt = rich_text(fb).strip()
        if fb_txt:
            lineas.append(fb)
    else:
        fb_str = (fb or "").strip()
        if fb_str:
            lineas.append(rich_obj(fb_str, [{"text": fb_str}]))
    return lineas


IMAGE_LABEL_RED = "C00000"

def add_image_label_paragraph(texto: str, style: str) -> str:
    texto = texto.strip()
    if re.match(r"^Pie de imagen:\s*", texto, re.I):
        resto = re.sub(r"^(Pie de imagen:\s*)+", "", texto, flags=re.I)
        return (
            f'    <w:p><w:pPr><w:pStyle w:val="{style}"/></w:pPr>'
            f'<w:r><w:rPr><w:color w:val="{IMAGE_LABEL_RED}"/></w:rPr>'
            f'<w:t xml:space="preserve">Pie de imagen: </w:t></w:r>'
            f'<w:r><w:t>{esc(resto)}</w:t></w:r>'
            f'</w:p>'
        )
    elif re.match(r"^Descripci[oó]n de (la )?imagen:\s*", texto, re.I):
        m = re.match(r"^(Descripci[oó]n de (la )?imagen:)\s*", texto, re.I)
        label = m.group(1)
        resto = re.sub(r"^(Descripci[oó]n de (la )?imagen:\s*)+", "", texto, flags=re.I)
        return (
            f'    <w:p><w:pPr><w:pStyle w:val="{style}"/></w:pPr>'
            f'<w:r><w:rPr><w:color w:val="{IMAGE_LABEL_RED}"/></w:rPr>'
            f'<w:t xml:space="preserve">{esc(label)} </w:t></w:r>'
            f'<w:r><w:t>{esc(resto)}</w:t></w:r>'
            f'</w:p>'
        )
    elif re.match(r"^Imagen_\d+$", texto, re.I):
        return (
            f'    <w:p><w:pPr><w:pStyle w:val="{style}"/></w:pPr>'
            f'<w:r><w:rPr><w:color w:val="{IMAGE_LABEL_RED}"/></w:rPr>'
            f'<w:t>{esc(texto)}</w:t></w:r>'
            f'</w:p>'
        )
    else:
        return p(texto, style)

def p_pie_imagen(texto: str) -> str:
    texto = texto.strip()
    if not re.match(r"^Pie de imagen:", texto, re.I):
        texto = "Pie de imagen: " + texto
    return add_image_label_paragraph(texto, "Cuerpoparrafo")


def p_desc_imagen(texto: str) -> str:
    texto = texto.strip()
    if not re.match(r"^Descripci[oó]n de (la )?imagen:", texto, re.I):
        texto = "Descripción de la imagen: " + texto
    return add_image_label_paragraph(texto, "Normal")





# =============================================================================
# Rasterización segura de esquemas/SmartArt
# =============================================================================

def _mime_img(ext: str) -> str:
    ext = (ext or '').lower().lstrip('.')
    return {
        'png': 'image/png',
        'jpg': 'image/jpeg',
        'jpeg': 'image/jpeg',
        'gif': 'image/gif',
        'bmp': 'image/bmp',
        'tif': 'image/tiff',
        'tiff': 'image/tiff',
    }.get(ext, 'image/png')


def _xml_img_raster(rid: str, cx: int, cy: int, docpr_id: int, nombre: str = 'Esquema') -> str:
    """Devuelve un párrafo WordprocessingML con una imagen raster centrada."""
    cx = int(cx or 5400040)
    cy = int(cy or 3150235)
    return (
        f'    <w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:drawing>'
        f'<wp:inline distT="0" distB="0" distL="0" distR="0">'
        f'<wp:extent cx="{cx}" cy="{cy}"/><wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:docPr id="{docpr_id}" name="{esc(nombre)} {docpr_id}"/><wp:cNvGraphicFramePr/>'
        f'<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        f'<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'<pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'<pic:nvPicPr><pic:cNvPr id="0" name="{esc(nombre)}.png"/><pic:cNvPicPr/></pic:nvPicPr>'
        f'<pic:blipFill><a:blip r:embed="{rid}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
        f'<pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr>'
        f'</pic:pic></a:graphicData></a:graphic></wp:inline></w:drawing></w:r></w:p>'
    )


def _exportar_imagenes_html_docx(docx_path: Path) -> list[dict]:
    """
    LibreOffice exporta imágenes, SmartArt y dibujos como raster en HTML.
    Usamos esa salida como captura fiel de los esquemas para evitar que Word
    recalcule SmartArt y deforme cajas, flechas o textos al reconstruir el DOCX.
    """
    if not docx_path or not docx_path.exists() or not _es_zip(docx_path):
        return []

    try:
        with tempfile.TemporaryDirectory(prefix='conv_papel_html_') as td:
            outdir = Path(td)
            cmd = [
                'libreoffice', '--headless', '--convert-to', 'html',
                '--outdir', str(outdir), str(docx_path)
            ]
            subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=90, check=False)
            htmls = list(outdir.glob('*.html'))
            if not htmls:
                return []
            html = htmls[0].read_text(encoding='utf-8', errors='ignore')
            imgs: list[dict] = []
            for m in re.finditer(r'<img\s+[^>]*src="([^"]+)"[^>]*>', html, flags=re.I):
                tag = m.group(0)
                src = m.group(1)
                img_path = outdir / src
                if not img_path.exists():
                    continue
                wm = re.search(r'\bwidth="(\d+)"', tag)
                hm = re.search(r'\bheight="(\d+)"', tag)
                width = int(wm.group(1)) if wm else None
                height = int(hm.group(1)) if hm else None
                data = img_path.read_bytes()
                ext = img_path.suffix.lower().lstrip('.') or 'png'

                # Normalizamos a PNG; evita problemas de GIF de SmartArt en algunos Word/LibreOffice.
                try:
                    from PIL import Image
                    import io
                    im = Image.open(img_path)
                    if im.mode not in ('RGB', 'RGBA'):
                        im = im.convert('RGBA' if 'A' in im.getbands() else 'RGB')
                    bio = io.BytesIO()
                    im.save(bio, format='PNG')
                    data = bio.getvalue()
                    ext = 'png'
                    if not width or not height:
                        width, height = im.size
                except Exception:
                    pass

                imgs.append({'bytes': data, 'ext': ext, 'width': width, 'height': height, 'src': src})
            return imgs
    except Exception:
        return []


def _extent_parrafo_grafico(par_xml: str) -> tuple[int, int]:
    m = re.search(r'<wp:extent\s+cx="(\d+)"\s+cy="(\d+)"', par_xml or '')
    if m:
        return int(m.group(1)), int(m.group(2))
    return 5400040, 3150235


def _preparar_rasteres_graficos(graficos: list[dict], docx_path: Path) -> None:
    """Asocia a cada gráfico extraído su captura rasterizada por orden de aparición."""
    imgs = _exportar_imagenes_html_docx(docx_path)
    if not imgs:
        return
    for g in graficos:
        ord_ = g.get('ordinal_grafico')
        if isinstance(ord_, int) and 0 <= ord_ < len(imgs):
            img = imgs[ord_]
            g['raster_bytes'] = img['bytes']
            g['raster_ext'] = img.get('ext', 'png')
            g['raster_width'] = img.get('width')
            g['raster_height'] = img.get('height')


def _aplicar_raster_a_graficos(graficos: list[dict]) -> None:
    """Sustituye el XML SmartArt/dibujo por una imagen estática fiel al original."""
    seq = 1
    for g in graficos:
        data = g.get('raster_bytes')
        if not data:
            continue
        rid = f'rIdRaster{seq}'
        media_name = f'esquema_raster_{seq}.{g.get("raster_ext", "png")}'
        cx, cy = _extent_parrafo_grafico(g.get('xml', ''))
        max_cx = 6200000
        if cx > max_cx:
            ratio = max_cx / cx
            cx = int(cx * ratio)
            cy = int(cy * ratio)
        g['raster_rid'] = rid
        g['raster_media'] = media_name
        g['xml'] = _xml_img_raster(rid, cx, cy, 9000 + seq, 'Esquema')
        seq += 1


def _registrar_rasteres_en_paquete(archivos: dict[str, bytes], graficos: list[dict]) -> None:
    """Añade las imágenes raster al paquete DOCX y sus relaciones."""
    rasteres = [g for g in graficos if g.get('raster_bytes') and g.get('raster_rid') and g.get('raster_media')]
    if not rasteres:
        return

    for g in rasteres:
        archivos[f"word/media/{g['raster_media']}"] = g['raster_bytes']

    ct = archivos.get('[Content_Types].xml', b'').decode('utf-8', errors='ignore')
    if ct and '</Types>' in ct:
        existing_exts = set(re.findall(r'<Default\s+[^>]*Extension="([^"]+)"', ct))
        inserts = []
        for g in rasteres:
            ext = str(g.get('raster_ext', 'png')).lower().lstrip('.') or 'png'
            if ext not in existing_exts:
                inserts.append(f'<Default Extension="{esc(ext)}" ContentType="{_mime_img(ext)}"/>')
                existing_exts.add(ext)
        if inserts:
            archivos['[Content_Types].xml'] = ct.replace('</Types>', ''.join(inserts) + '</Types>').encode('utf-8')

    key = 'word/_rels/document.xml.rels'
    rels = _parse_relationships_xml(archivos.get(key, b''))
    by_id = {r['Id']: r for r in rels}
    for g in rasteres:
        rid = g['raster_rid']
        if rid in by_id:
            continue
        rels.append({
            'Id': rid,
            'Type': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image',
            'Target': f"media/{g['raster_media']}"
        })
    archivos[key] = _rels_to_xml(rels)


# =============================================================================
# Interacciones
# =============================================================================

def parsear_interacciones(path: Path | None) -> dict[int, dict]:
    if not path or not path.exists():
        return {}

    if not _es_zip(path):
        return parsear_interacciones_texto(path)

    from docx import Document
    from docx.oxml.ns import qn as _qn

    def _cell_lineas(cell) -> list[str]:
        """
        Extrae las líneas de una celda párrafo a párrafo.
        Cuando un párrafo tiene numeración automática (w:numPr) asigna el
        prefijo de letra correspondiente (a., b., …) para no perder ese dato
        al leer solo cell.text.
        """
        letras_seq = "abcdefghijklmnopqrstuvwxyz"
        letra_idx: dict[str, int] = {}  # numId_ilvl → contador
        lines: list[str] = []
        for para in cell.paragraphs:
            txt = _norm_line(para.text)
            if not txt:
                continue
            # Detectar numeración automática
            try:
                pPr = para._p.find(_qn("w:pPr"))
                numPr = pPr.find(_qn("w:numPr")) if pPr is not None else None
                if numPr is not None:
                    numId_el = numPr.find(_qn("w:numId"))
                    ilvl_el = numPr.find(_qn("w:ilvl"))
                    numId_val = int(numId_el.get(_qn("w:val"), "0")) if numId_el is not None else 0
                    ilvl_val = int(ilvl_el.get(_qn("w:val"), "0")) if ilvl_el is not None else 0
                    if numId_val > 0:
                        _num_fmt = _inter_num_fmt.get((str(numId_val), str(ilvl_val)), "")
                        key = f"{numId_val}_{ilvl_val}"
                        if _num_fmt == "bullet":
                            _vsim = {0: "●", 1: "○", 2: "▪", 3: "–"}
                            prefix = _vsim.get(ilvl_val, "●") + " "
                            clean_txt = re.sub(r"^[a-zA-Z0-9]+[.)]\s+", "", txt)
                            txt = prefix + clean_txt
                        else:
                            idx = letra_idx.get(key, 0)
                            if idx < len(letras_seq):
                                txt = letras_seq[idx] + ". " + txt
                            letra_idx[key] = idx + 1
            except Exception:
                pass
            lines.append(txt)
        return lines

    def _cell_lineas_rich(cell) -> list[dict]:
        """Like _cell_lineas but returns rich_obj per paragraph, preserving bold/italic runs."""
        letras_seq = "abcdefghijklmnopqrstuvwxyz"
        letra_idx: dict[str, int] = {}
        rich_lines = []
        for para in cell.paragraphs:
            txt = _norm_line(para.text)
            if not txt:
                continue
            runs = []
            for r_el in para._p.findall('.//' + _qn('w:r')):
                texts = [t.text or '' for t in r_el.findall('.//' + _qn('w:t'))]
                text = ''.join(texts)
                if not text:
                    continue
                rpr = r_el.find(_qn('w:rPr'))
                bold_r = rpr is not None and rpr.find(_qn('w:b')) is not None
                italic_r = rpr is not None and rpr.find(_qn('w:i')) is not None
                r: dict = {"text": text}
                if bold_r:
                    r["bold"] = True
                if italic_r:
                    r["italic"] = True
                runs.append(r)
            if not runs:
                runs = [{"text": txt}]
            try:
                pPr = para._p.find(_qn("w:pPr"))
                numPr = pPr.find(_qn("w:numPr")) if pPr is not None else None
                if numPr is not None:
                    numId_el = numPr.find(_qn("w:numId"))
                    ilvl_el = numPr.find(_qn("w:ilvl"))
                    numId_val = int(numId_el.get(_qn("w:val"), "0")) if numId_el is not None else 0
                    ilvl_val = int(ilvl_el.get(_qn("w:val"), "0")) if ilvl_el is not None else 0
                    if numId_val > 0:
                        _num_fmt = _inter_num_fmt.get((str(numId_val), str(ilvl_val)), "")
                        key = f"{numId_val}_{ilvl_val}"
                        if _num_fmt == "bullet":
                            _vsim = {0: "●", 1: "○", 2: "▪", 3: "–"}
                            prefix = _vsim.get(ilvl_val, "●") + " "
                            # Strip any manually-typed letter/number prefix (e.g. "a. ", "1. ")
                            mod_runs = [dict(r) for r in runs]
                            if mod_runs:
                                first_clean = re.sub(r"^[a-zA-Z0-9]+[.)]\s+", "", mod_runs[0].get("text", ""))
                                if first_clean != mod_runs[0].get("text", ""):
                                    mod_runs[0]["text"] = first_clean
                                    if not mod_runs[0]["text"]:
                                        mod_runs = mod_runs[1:]
                            clean_txt = re.sub(r"^[a-zA-Z0-9]+[.)]\s+", "", txt)
                            runs = [{"text": prefix}] + mod_runs
                            txt = prefix + clean_txt
                        else:
                            idx = letra_idx.get(key, 0)
                            if idx < len(letras_seq):
                                prefix = letras_seq[idx] + ". "
                                runs = [{"text": prefix}] + runs
                                txt = prefix + txt
                            letra_idx[key] = idx + 1
            except Exception:
                pass
            rich_lines.append(rich_obj(txt, runs))
        return rich_lines

    def _opciones_rich_data(rich_lines: list[dict]) -> dict:
        """Zone rich_lines into opciones_rich, solucion_rich, feedback_rich."""
        opciones_rich: list[dict] = []
        solucion_rich = None
        feedback_rich = None
        zona = "opciones"
        visto_letra = False

        for ro in rich_lines:
            txt = rich_text(ro).strip()
            if not txt:
                continue
            if re.match(r"^Opciones:?$", txt, re.I):
                zona = "opciones"
                continue
            if re.match(r"^Soluci[oó]n:", txt, re.I):
                zona = "solucion"
                rest_ro = _strip_prefix_runs(ro, r"^Soluci[oó]n:\s*")
                if rich_text(rest_ro).strip():
                    solucion_rich = rest_ro
                continue
            if re.match(r"^(Feedback|Retroalimentaci[oó]n):", txt, re.I):
                zona = "feedback"
                rest_ro = _strip_prefix_runs(ro, r"^(Feedback|Retroalimentaci[oó]n):\s*")
                if rich_text(rest_ro).strip():
                    feedback_rich = rest_ro
                continue
            if zona == "opciones":
                if re.match(r"^Actividad de evaluaci[oó]n$", txt, re.I):
                    continue
                if re.match(r"^[a-h][\).]\s*", txt):
                    visto_letra = True
                    opciones_rich.append(_strip_prefix_runs(ro, r"^[a-h][\).]\s*"))
                elif not visto_letra:
                    pass  # enunciado already in plain parse
                else:
                    if opciones_rich:
                        last = opciones_rich[-1]
                        combined_txt = rich_text(last) + " " + txt
                        combined_runs = rich_runs(last) + [{"text": " "}] + rich_runs(ro)
                        opciones_rich[-1] = rich_obj(combined_txt, combined_runs)
            elif zona == "solucion":
                if solucion_rich is None:
                    solucion_rich = ro
                else:
                    combined_txt = rich_text(solucion_rich) + " " + txt
                    combined_runs = rich_runs(solucion_rich) + [{"text": " "}] + rich_runs(ro)
                    solucion_rich = rich_obj(combined_txt, combined_runs)
            elif zona == "feedback":
                if feedback_rich is None:
                    feedback_rich = ro
                else:
                    combined_txt = rich_text(feedback_rich) + " " + txt
                    combined_runs = rich_runs(feedback_rich) + [{"text": " "}] + rich_runs(ro)
                    feedback_rich = rich_obj(combined_txt, combined_runs)

        return {
            "opciones_rich": opciones_rich,
            "solucion_rich": solucion_rich,
            "feedback_rich": feedback_rich,
        }

    doc = Document(str(path))
    result: dict[int, dict] = {}

    _inter_num_fmt: dict[tuple, str] = {}
    try:
        nbp = doc.part.numbering_part
        if nbp is not None:
            nxml = nbp._element
            _abs_defs: dict = {}
            for _abn in nxml.findall(_qn("w:abstractNum")):
                _abs_id = _abn.get(_qn("w:abstractNumId"))
                _levels: dict = {}
                for _lvl in _abn.findall(_qn("w:lvl")):
                    _ilvl = _lvl.get(_qn("w:ilvl"))
                    _fmt_el = _lvl.find(_qn("w:numFmt"))
                    if _fmt_el is not None:
                        _levels[_ilvl] = _fmt_el.get(_qn("w:val"), "")
                _abs_defs[_abs_id] = _levels
            for _num in nxml.findall(_qn("w:num")):
                _num_id = _num.get(_qn("w:numId"))
                _abs_ref = _num.find(_qn("w:abstractNumId"))
                if _abs_ref is not None:
                    _abs_id = _abs_ref.get(_qn("w:val"))
                    if _abs_id in _abs_defs:
                        for _ilvl, _fmt in _abs_defs[_abs_id].items():
                            _inter_num_fmt[(_num_id, _ilvl)] = _fmt
    except Exception:
        pass

    for tbl in doc.tables:
        if not tbl.rows:
            continue

        header = tbl.rows[0].cells[0].text.strip() if tbl.rows[0].cells else ""
        m = RE_INTER.match(header)
        if not m:
            continue

        n = int(m.group(1))
        if len(tbl.rows) > 1 and tbl.rows[1].cells:
            cell = tbl.rows[1].cells[0]
            lines_raw = _cell_lineas(cell)
            raw = "\n".join(lines_raw)
        else:
            raw = ""
            cell = None
        raw = raw.replace("\r", "\n").strip()

        _tiene_opciones = bool(
            re.search(r"(?im)^\s*Opciones:?\s*$", raw)
            or re.search(r"(?m)^[a-d][.)]\s+\S", raw)
        )
        if _tiene_opciones and re.search(r"(?im)^\s*Soluci[oó]n:", raw):
            result[n] = _parsear_interaccion_opciones(raw)
            if cell is not None:
                rich_lines = _cell_lineas_rich(cell)
                result[n].update(_opciones_rich_data(rich_lines))
        elif re.search(r"(?im)^\s*Desplegables:?\s*$", raw):
            result[n] = _parsear_interaccion_desplegables(raw)
            if cell is not None:
                rich_lines = _cell_lineas_rich(cell)
                result[n]["items_rich"] = _parsear_interaccion_desplegables_rich(rich_lines)
        else:
            result[n] = {
                "tipo": "texto",
                "lineas": [_norm_line(x) for x in raw.splitlines() if _norm_line(x)]
            }

    return result


def parsear_interacciones_texto(path: Path) -> dict[int, dict]:
    text = path.read_text(encoding="utf-8", errors="replace")
    result: dict[int, dict] = {}

    parts = re.split(r"(?=Interacci[oó]n\s+\d+)", text)

    for part in parts:
        part = part.strip()
        if not part:
            continue

        lines = [re.sub(r"^\|?|\|?$", "", x).strip() for x in part.splitlines()]
        lines = [
            re.sub(r"\*+", "", x).strip()
            for x in lines
            if x.strip() and not re.match(r"^\|?\s*-+\s*\|?$", x)
        ]

        if not lines:
            continue

        m = RE_INTER.match(lines[0])
        if not m:
            continue

        n = int(m.group(1))
        raw = "\n".join(lines[1:]).strip()

        _tiene_opciones = bool(
            re.search(r"(?im)^\s*Opciones:?\s*$", raw)
            or re.search(r"(?m)^[a-d][.)]\s+\S", raw)
        )
        if _tiene_opciones and re.search(r"(?im)^\s*Soluci[oó]n:", raw):
            result[n] = _parsear_interaccion_opciones(raw)
        elif re.search(r"(?im)^\s*Desplegables:?\s*$", raw):
            result[n] = _parsear_interaccion_desplegables(raw)

    return result


def _parsear_interaccion_opciones(raw: str) -> dict:
    opciones: list[str] = []
    enunciado_lines: list[str] = []
    solucion: list[str] = []
    feedback: list[str] = []
    zona = "opciones"
    visto_letra = False  # True cuando ya se vio la primera opción con letra (a), b)…)

    for line in raw.splitlines():
        line = line.strip()

        if not line:
            continue

        if re.match(r"^Opciones:?$", line, re.I):
            zona = "opciones"
            continue

        if re.match(r"^Soluci[oó]n:", line, re.I):
            zona = "solucion"
            rest = line.split(":", 1)[1].strip() if ":" in line else ""
            if rest:
                solucion.append(rest)
            continue

        if re.match(r"^(Feedback|Retroalimentaci[oó]n):", line, re.I):
            zona = "feedback"
            rest = line.split(":", 1)[1].strip() if ":" in line else ""
            if rest:
                feedback.append(rest)
            continue

        if zona == "opciones":
            if re.match(r"^Actividad de evaluaci[oó]n$", line, re.I):
                continue
            if re.match(r"^[a-h][\).]\s*", line):
                visto_letra = True
                clean = re.sub(r"^[a-h][\).]\s*", "", line)
                opciones.append(clean)
            elif not visto_letra:
                # Texto antes de la primera opción con letra → es el enunciado
                enunciado_lines.append(line)
            else:
                # Continuación de la última opción
                if opciones:
                    opciones[-1] += " " + line
        elif zona == "solucion":
            solucion.append(line)
        elif zona == "feedback":
            feedback.append(line)

    return {
        "tipo": "opciones",
        "enunciado": " ".join(enunciado_lines).strip(),
        "opciones": opciones,
        "solucion": " ".join(solucion).strip(),
        "feedback": " ".join(feedback).strip(),
    }


def _parsear_interaccion_desplegables(raw: str) -> dict:
    # Acepta "Desplegables:" y también "Desplegables" sin dos puntos.
    m = re.search(r"(?im)^\s*Desplegables:?\s*$", raw)
    if m:
        raw_items = raw[m.end():]
    elif "Desplegables:" in raw:
        raw_items = raw.split("Desplegables:", 1)[1]
    else:
        raw_items = raw
    lines = [x.rstrip() for x in raw_items.splitlines()]

    items = []
    i = 0

    while i < len(lines):
        while i < len(lines) and not lines[i].strip():
            i += 1

        if i >= len(lines):
            break

        titulo = lines[i].strip()
        i += 1
        contenido: list[str] = []

        while i < len(lines):
            current = lines[i].strip()

            if current and contenido and _parece_titulo_desplegable(current):
                nxt = _siguiente_no_vacia(lines, i + 1)
                if nxt is not None:
                    break

            contenido.append(lines[i])
            i += 1

        body, subitems = _separar_cuerpo_y_subitems(contenido)

        items.append({
            "titulo": titulo,
            "body": body,
            "subitems": subitems,
        })

    return {"tipo": "desplegables", "items": items}


def _parece_titulo_desplegable(line: str) -> bool:
    line = line.strip()

    if not line:
        return False
    if len(line) > 90:
        return False
    if line[-1] in ".;:":
        return False
    if line[0].islower():
        return False
    if re.match(r"^[·●○▪–\-]\s", line):
        return False
    if re.match(
        r"^(Por ejemplo|Asimismo|Además|También|Puede|Se calcula|Número|Cuando|Una|Un consumidor|El cliente|La empresa)",
        line
    ):
        return False

    return True


def _siguiente_no_vacia(lines: list[str], start: int) -> str | None:
    for j in range(start, len(lines)):
        if lines[j].strip():
            return lines[j].strip()
    return None


def _separar_cuerpo_y_subitems(lines: list[str]) -> tuple[str, list[str]]:
    clean = [x.strip() for x in lines]

    while clean and not clean[0]:
        clean.pop(0)
    while clean and not clean[-1]:
        clean.pop()

    if not clean:
        return "", []

    # Si el contenido del desplegable empieza con una viñeta del online,
    # no lo convertimos en el cuerpo del título. Debe conservarse como
    # sublista; de lo contrario, el primer punto queda pegado al título
    # principal y los recursos de imagen pasan a viñetas.
    if re.match(r"^[·●○▪\-–]\s*", clean[0]):
        return "", [x for x in clean if x.strip()]

    # Si el desplegable contiene recursos de imagen, todo el contenido debe
    # conservarse como secuencia interna: subpuntos + URL + pie + descripción.
    # Así evitamos que el primer subpunto se fusione con el título principal
    # y que las URL/pies/descripciones se conviertan en viñetas.
    if any(RE_URL.match(x) or re.match(r"^Pie de imagen:", x, re.I) or re.match(r"^Descripci[oó]n de (la )?imagen:", x, re.I) for x in clean):
        return "", [x for x in clean if x.strip()]

    # Si hay varias líneas tipo "Etiqueta: explicación", son subpuntos del
    # desplegable, no un único párrafo unido.
    colon_lines = [x for x in clean if x.strip()]
    if len(colon_lines) > 1 and all(":" in x and len(x.split(":", 1)[0]) <= 70 for x in colon_lines):
        return "", colon_lines

    bullet_idx = -1
    for k, line in enumerate(clean):
        if re.match(r"^[·●○▪\-–]\s*|^o\s+", line, re.I):
            bullet_idx = k
            break

    blank_idx = clean.index("") if "" in clean else -1

    split_idx = -1
    if blank_idx != -1 and bullet_idx != -1:
        split_idx = min(blank_idx, bullet_idx)
    elif blank_idx != -1:
        split_idx = blank_idx
    elif bullet_idx != -1:
        split_idx = bullet_idx

    if split_idx == -1:
        return _join_sin_doble_puntuacion(clean), []

    body_lines = [x for x in clean[:split_idx] if x.strip()]
    rest = [x for x in clean[split_idx:] if x.strip()]

    return _join_sin_doble_puntuacion(body_lines), rest


def _join_sin_doble_puntuacion(lines: list[str]) -> str:
    text = " ".join(x.strip() for x in lines if x.strip())
    text = re.sub(r"::+", ":", text)
    text = re.sub(r"\s+([,.;:])", r"\1", text)
    return text.strip()


def _strip_prefix_runs(ro: dict, prefix_re: str) -> dict:
    """Strip a regex prefix from the text/runs of a rich_obj."""
    txt = rich_text(ro)
    m = re.match(prefix_re, txt, re.I)
    if not m:
        return ro
    skip = m.end()
    runs = [dict(r) for r in rich_runs(ro)]
    consumed = 0
    new_runs = []
    for r in runs:
        rt = str(r.get("text", ""))
        if consumed < skip:
            remaining = skip - consumed
            if len(rt) <= remaining:
                consumed += len(rt)
                continue
            r["text"] = rt[remaining:]
            consumed = skip
        new_runs.append(r)
    new_txt = txt[skip:]
    return rich_obj(new_txt, new_runs or [{"text": new_txt}])


def _join_rich_objs(pairs: list[tuple]) -> dict:
    """Join (rich_obj, text) pairs into a single rich_obj with space between paragraphs."""
    if not pairs:
        return rich_obj("", [])
    all_runs: list[dict] = []
    combined_text = ""
    for k, (ro, t) in enumerate(pairs):
        if k > 0 and t:
            all_runs.append({"text": " "})
            combined_text += " "
        all_runs.extend(rich_runs(ro))
        combined_text += t
    return rich_obj(combined_text, all_runs or [{"text": combined_text}])


def _separar_cuerpo_y_subitems_rich(rich_items: list[dict]) -> tuple:
    """Rich version of _separar_cuerpo_y_subitems. Returns (body_rich|None, subitems_rich)."""
    texts = [rich_text(ro).strip() for ro in rich_items]
    while rich_items and not texts[0]:
        rich_items, texts = rich_items[1:], texts[1:]
    while rich_items and not texts[-1]:
        rich_items, texts = rich_items[:-1], texts[:-1]
    if not rich_items:
        return None, []
    if re.match(r"^[·●○▪\-–]\s*", texts[0]):
        return None, [ro for ro, t in zip(rich_items, texts) if t]
    if any(RE_URL.match(t) or re.match(r"^Pie de imagen:", t, re.I)
           or re.match(r"^Descripci[oó]n de (la )?imagen:", t, re.I) for t in texts if t):
        return None, [ro for ro, t in zip(rich_items, texts) if t]
    non_empty = [(ro, t) for ro, t in zip(rich_items, texts) if t]
    if len(non_empty) > 1 and all(":" in t and len(t.split(":", 1)[0]) <= 70 for _, t in non_empty):
        return None, [ro for ro, _ in non_empty]
    bullet_idx = -1
    for k, t in enumerate(texts):
        if re.match(r"^[·●○▪\-–]\s*|^o\s+", t, re.I):
            bullet_idx = k
            break

    blank_idx = texts.index("") if "" in texts else -1

    split_idx = -1
    if blank_idx != -1 and bullet_idx != -1:
        split_idx = min(blank_idx, bullet_idx)
    elif blank_idx != -1:
        split_idx = blank_idx
    elif bullet_idx != -1:
        split_idx = bullet_idx

    if split_idx == -1:
        return _join_rich_objs(non_empty), []

    body_pairs = [(ro, t) for ro, t in zip(rich_items[:split_idx], texts[:split_idx]) if t]
    rest_pairs = [(ro, t) for ro, t in zip(rich_items[split_idx:], texts[split_idx:]) if t]
    return (_join_rich_objs(body_pairs) if body_pairs else None), [ro for ro, _ in rest_pairs]


def _parsear_interaccion_desplegables_rich(rich_lines: list[dict]) -> list[dict]:
    """Returns items_rich: [{titulo, titulo_rich, body_rich, subitems_rich}]."""
    start = 0
    for i, ro in enumerate(rich_lines):
        if re.match(r"^Desplegables:?$", rich_text(ro).strip(), re.I):
            start = i + 1
            break
    rlines = rich_lines[start:]
    items_rich: list[dict] = []
    i, n = 0, len(rlines)
    while i < n:
        while i < n and not rich_text(rlines[i]).strip():
            i += 1
        if i >= n:
            break
        titulo_ro = rlines[i]
        titulo_txt = rich_text(titulo_ro).strip()
        i += 1
        contenido: list[dict] = []
        while i < n:
            cro = rlines[i]
            ctxt = rich_text(cro).strip()
            if ctxt and contenido and _parece_titulo_desplegable(ctxt):
                j = i + 1
                while j < n and not rich_text(rlines[j]).strip():
                    j += 1
                if j < n:
                    break
            contenido.append(cro)
            i += 1
        body_rich, subitems_rich = _separar_cuerpo_y_subitems_rich(list(contenido))
        items_rich.append({
            "titulo": titulo_txt,
            "titulo_rich": titulo_ro,
            "body_rich": body_rich,
            "subitems_rich": subitems_rich,
        })
    return items_rich


def expandir_interaccion(n: int, interacciones: dict[int, dict], label: str = "") -> list[dict]:
    inter = interacciones.get(n)

    if not inter:
        return []
    if inter.get("tipo") == "opciones":
        return []

    # Interacciones de tipo texto: renderizar según el label del marcador
    if inter.get("tipo") == "texto":
        lineas_raw = inter.get("lineas", [])
        lineas_rich = []
        extra_blocks: list[dict] = []
        for linea in lineas_raw:
            if not linea:
                continue
            if debe_elim(linea):
                continue
            if RE_URL.match(linea):
                extra_blocks.append({
                    "tipo": "url_imagen" if _es_url_imagen(linea) else "url_recurso",
                    "texto": linea,
                })
            elif linea.startswith("Pie de imagen:"):
                extra_blocks.append({"tipo": "pie_imagen", "texto": linea[len("Pie de imagen:"):].strip()})
            elif re.match(r"^Descripci[oó]n de (la )?imagen:", linea):
                extra_blocks.append({
                    "tipo": "desc_imagen",
                    "texto": re.sub(r"^Descripci[oó]n de (la )?imagen:\s*", "", linea),
                })
            else:
                lineas_rich.append(rich_obj(linea, [{"text": linea}]))

        # Detectar tipo de actividad por el label
        mt_colab = re.search(r"Actividad colaborativa\s*(\d*)", label, re.I)
        mt_eval = re.search(r"Actividad de evaluaci[oó]n", label, re.I)
        mt_apr = re.search(r"Actividad de aprendizaje\s+(\d+)", label, re.I)
        mt_aprac = re.search(r"Aplicaci[oó]n pr[aá]ctica\s*(\d*)", label, re.I)
        mt_tarea = re.search(r"Tarea de evaluaci[oó]n\s*(\d*)", label, re.I)

        if mt_colab:
            num = mt_colab.group(1).strip()
            return [{"tipo": "actividad_complementaria", "etiqueta": "Actividad complementaria",
                     "lineas": lineas_rich, "_num_collab": num}] + extra_blocks
        elif mt_eval and lineas_rich:
            num_eval = re.search(r"(\d+)", label[mt_eval.end():]) or re.search(r"(\d+)", label)
            etiqueta = f"Actividad {num_eval.group(1)}" if num_eval else f"Actividad {n}"
            return [{"tipo": "tarea", "etiqueta": etiqueta, "lineas": lineas_rich}] + extra_blocks
        elif mt_apr and lineas_rich:
            num = mt_apr.group(1).strip()
            etiqueta = f"Aplicación práctica {num}".strip() if num else "Aplicación práctica"
            return [{"tipo": "aplicacion_practica", "etiqueta": etiqueta, "lineas": lineas_rich}] + extra_blocks
        elif mt_aprac and lineas_rich:
            num = mt_aprac.group(1).strip()
            etiqueta = f"Tarea {num}".strip() if num else "Tarea"
            return [{"tipo": "tarea", "etiqueta": etiqueta, "lineas": lineas_rich}] + extra_blocks
        elif mt_tarea and lineas_rich:
            num = mt_tarea.group(1).strip()
            etiqueta = f"Tarea {num}".strip() if num else "Tarea"
            return [{"tipo": "tarea", "etiqueta": etiqueta, "lineas": lineas_rich}] + extra_blocks
        else:
            # Sin label especial: líneas como párrafos sueltos
            bloques: list[dict] = []
            for ro in lineas_rich:
                txt = rich_text(ro)
                limpio = _limpiar_vineta_literal(txt)
                if limpio:
                    if _abre_modo_lista(txt):
                        bloques.append({"tipo": "p_vineta", "texto": rich_obj(limpio, [{"text": limpio}]), "nivel": 1})
                    else:
                        bloques.append({"tipo": "parrafo", "texto": ro})
            return bloques + extra_blocks

    if inter.get("tipo") != "desplegables":
        return []

    bloques: list[dict] = []

    items_rich = inter.get("items_rich", [])
    rich_map = {item["titulo"]: item for item in items_rich}

    for item in inter.get("items", []):
        titulo = item.get("titulo", "").strip()
        body = item.get("body", "").strip()
        subitems = item.get("subitems", [])

        rich_item = rich_map.get(titulo, {})
        body_rich = rich_item.get("body_rich")
        subitems_rich = rich_item.get("subitems_rich", [])

        if not titulo or debe_elim(titulo):
            continue

        if body and subitems:
            bloques.append({
                "tipo": "desplegable_procedimental",
                "titulo": titulo,
                "titulo_rich": rich_item.get("titulo_rich"),
                "body": body,
                "body_rich": body_rich,
                "subitems": subitems,
                "subitems_rich": subitems_rich,
            })
            continue

        if body:
            bloques.append({
                "tipo": "desplegable_simple",
                "titulo": titulo,
                "titulo_rich": rich_item.get("titulo_rich"),
                "contenido": body,
                "contenido_rich": body_rich,
            })
            continue

        # No body: decide inline vs two-level based on whether subitems have "term: desc" format.
        # If the first subitem has a colon split (term: description), render as:
        #   ● Title:          (bold bullet, level 1)
        #   ○ Term: desc...   (p_desp inline, level 2)
        # If the first subitem is plain text, absorb it inline:
        #   ● Title: description... (p_desp inline, level 1)
        sub_start = 0
        if subitems:
            first_sub = _limpiar_vineta_literal(subitems[0])
            first_sub_rich_raw = subitems_rich[0] if subitems_rich else None
            first_has_colon = (
                ":" in first_sub
                and len(first_sub.split(":", 1)[0].strip()) <= 70
                and not re.match(r"^(Pie de imagen|Descripci[oó]n de (la )?imagen|https?://)", first_sub, re.I)
            )
            if first_sub and not first_has_colon:
                # Plain description → inline with title
                bloques.append({
                    "tipo": "desplegable_simple",
                    "titulo": titulo,
                    "titulo_rich": rich_item.get("titulo_rich"),
                    "contenido": first_sub,
                    "contenido_rich": _limpiar_vineta_rich(first_sub_rich_raw) if first_sub_rich_raw else None,
                })
                sub_start = 1
            else:
                # Term:desc subitems → title as bold bullet, all subitems at level 2
                tit_r = rich_item.get("titulo_rich")
                bloques.append({
                    "tipo": "p_vineta_bold",
                    "texto": _titulo_con_colon(tit_r) if tit_r else (titulo + ":"),
                    "nivel": 1,
                })
                sub_start = 0
        else:
            bloques.append({"tipo": "p_vineta_bold", "texto": titulo + ":", "nivel": 1})

        for j, sub in enumerate(subitems[sub_start:], sub_start):
            sub_rich = subitems_rich[j] if j < len(subitems_rich) else None
            sub = _limpiar_vineta_literal(sub)
            if not sub:
                continue

            if RE_URL.match(sub):
                bloques.append({
                    "tipo": "url_imagen" if _es_url_imagen(sub) else "url_recurso",
                    "texto": sub,
                })
            elif re.match(r"^Pie de imagen:", sub, re.I):
                bloques.append({"tipo": "pie_imagen", "texto": re.sub(r"^Pie de imagen:\s*", "", sub, flags=re.I)})
            elif re.match(r"^Descripci[oó]n de (la )?imagen:", sub, re.I):
                bloques.append({"tipo": "desc_imagen", "texto": re.sub(r"^Descripci[oó]n de (la )?imagen:\s*", "", sub, flags=re.I)})
            elif ":" in sub and len(sub.split(":", 1)[0]) <= 70:
                titulo_sub, desc_sub = sub.split(":", 1)
                desc_rich = None
                if sub_rich:
                    clean_rich = _limpiar_vineta_rich(sub_rich)
                    desc_rich = _strip_prefix_runs(clean_rich, re.escape(titulo_sub) + r"\s*:\s*")
                bloques.append({
                    "tipo": "desplegable_simple_n2",
                    "titulo": titulo_sub.strip(),
                    "contenido": desc_sub.strip(),
                    "contenido_rich": desc_rich,
                })
            else:
                bloques.append({
                    "tipo": "p_vineta",
                    "texto": sub_rich if sub_rich is not None else sub,
                    "nivel": 2,
                })

    return bloques


# =============================================================================
# Extracción PDF / TXT
# =============================================================================

def extraer_texto_pdf(pdf: Path) -> list[str]:
    from pypdf import PdfReader
    reader = PdfReader(str(pdf))
    return [page.extract_text() or "" for page in reader.pages]


def parsear_pdf_o_texto(path: Path, interacciones: dict | None = None) -> dict:
    if interacciones is None:
        interacciones = {}

    if path.suffix.lower() == ".pdf":
        paginas = extraer_texto_pdf(path)
        raw_lines = []
        for pag in paginas:
            raw_lines.extend(pag.splitlines())
    else:
        raw_lines = path.read_text(encoding="utf-8", errors="replace").splitlines()

    _tipo_map = {
        "Nota": "nota", "Ejemplo": "ejemplo",
        "Sabías que...": "sabias_que", "Sabías que…": "sabias_que",
        "Consejo": "consejo", "Definición": "definicion",
        "Hilo conductor": "hilo_conductor", "Para saber más": "para_saber_mas",
        "Vídeo": "video", "Video": "video",
        "Importante": "importante", "Recuerda": "recuerda",
        "Actividad complementaria": "actividad_complementaria",
        "Actividad colaborativa": "actividad_complementaria",
        "Actividad de evaluación": "tarea",
        "Actividad de aprendizaje": "aplicacion_practica",
        "Aplicación práctica": "tarea",
        "Caso práctico": "tarea", "Ejercicio": "tarea", "Tarea": "tarea",
    }

    est = {
        "titulo_unidad": "",
        "titulo_modulo": "",
        "objetivos": [],
        "secciones": [],
    }

    current_sec = None
    current_sub = None
    current_sub2 = None
    en_obj = False
    blk = None

    def activos() -> list:
        if current_sub2:
            return current_sub2["bloques"]
        if current_sub:
            return current_sub["bloques"]
        if current_sec:
            return current_sec["bloques"]
        return []

    _TIPOS_BLOQUE_INFO_PDF = {
        "nota", "ejemplo", "para_saber_mas", "video", "importante",
        "sabias_que", "consejo", "definicion", "hilo_conductor", "recuerda",
    }

    def _blk_info_activo_pdf():
        return blk is not None and blk.get("tipo") in _TIPOS_BLOQUE_INFO_PDF

    def flush_blk():
        nonlocal blk
        if not blk:
            return
        if blk.get("lineas") or blk.get("tipo") in {
            "hilo_conductor", "video", "ejemplo", "importante", "definicion", "sabias_que",
        }:
            b = {k: v for k, v in blk.items() if not k.startswith("_")}
            activos().append(b)
        blk = None

    def _add_inter_opciones(n, label, inter):
        mt_eval = re.search(r"Actividad de evaluaci[oó]n", label, re.I)
        mt_apr = re.search(r"Actividad de aprendizaje\s+(\d+)", label, re.I)
        mt_aprac = re.search(r"Aplicaci[oó]n pr[aá]ctica\s*(\d*)", label, re.I)

        if mt_eval:
            # "Actividad de evaluación X" (inline) → "Actividad X" sin solución
            num_eval = re.search(r"(\d+)", label[mt_eval.end():]) or re.search(r"(\d+)", label)
            etiqueta = f"Actividad {num_eval.group(1)}" if num_eval else f"Actividad {n}"
            tipo_final = "tarea"
        elif mt_apr:
            # "Actividad de aprendizaje X" (test) → "Aplicación práctica X" CON solución
            tipo_final = "aplicacion_practica"
            num_apr = mt_apr.group(1).strip()
            etiqueta = f"Aplicación práctica {num_apr}".strip() if num_apr else "Aplicación práctica"
        elif mt_aprac:
            tipo_final = "tarea"
            num_apr = mt_aprac.group(1).strip()
            etiqueta = f"Tarea {num_apr}".strip() if num_apr else "Tarea"
        else:
            tipo_final, etiqueta = "tarea", label.strip() or "Tarea"

        letras = "abcdefgh"
        opciones_rich = inter.get("opciones_rich", [])
        if opciones_rich:
            opciones = [{"letra": letras[j], "texto": opt_rich}
                        for j, opt_rich in enumerate(opciones_rich) if j < len(letras)]
        else:
            opciones = [{"letra": letras[j], "texto": opt.strip()}
                        for j, opt in enumerate(inter.get("opciones", []))]

        # El enunciado (texto de la pregunta) va como primera línea del bloque
        enunciado_text = inter.get("enunciado", "").strip()
        lineas = [rich_obj(enunciado_text, [{"text": enunciado_text}])] if enunciado_text else []

        entry = {"tipo": tipo_final, "etiqueta": etiqueta, "lineas": lineas, "opciones": opciones}
        if tipo_final == "aplicacion_practica":
            entry["solucion"] = inter.get("solucion_rich") or inter.get("solucion", "")
            entry["feedback"] = inter.get("feedback_rich") or inter.get("feedback", "")
        activos().append(entry)

    for raw in raw_lines:
        line = _norm_line(raw)
        if not line or es_solo_pua(line):
            continue

        # Unit / module title (before debe_elim)
        mtitulo = _match_titulo_unidad(line)
        if mtitulo:
            etiqueta_unidad, n_unidad, titulo_modulo = mtitulo
            est["titulo_unidad"] = f"{etiqueta_unidad} {n_unidad}"
            if titulo_modulo and not est["titulo_modulo"]:
                est["titulo_modulo"] = titulo_modulo
            continue

        if est["titulo_unidad"] and not est["titulo_modulo"] and not RE_SEC1.match(line) and not _es_cabecera_objetivos(line):
            if line != est["titulo_unidad"] and len(line) < 120 and not RE_URL.match(line):
                est["titulo_modulo"] = limpiar_titulo(line)
                continue

        # Objectives header
        if _es_cabecera_objetivos(line) and not current_sec:
            en_obj = True
            nueva_intro = _tipo_objetivos_intro(line)
            if "general" in line.lower():
                est["_current_obj_type"] = "general"
            else:
                est["_current_obj_type"] = "especifico"
            if "específico" in line.lower() or "especifico" in line.lower():
                est["_objectives_intro"] = nueva_intro
            continue

        # Inside objectives section
        if en_obj:
            if RE_SEC1.match(line) or line == "Introducción" or _es_cabecera_no_contenido(line) or any(line.startswith(p) for p in BLOQUES_ESP):
                en_obj = False
                # fall through to section handling below
            else:
                if debe_elim(line):
                    continue
                if est.get("_current_obj_type") == "general":
                    est["objetivo_general"] = line
                    continue
                limpio = _limpiar_vineta_literal(line)
                if limpio and not re.match(r"^CE\s+[a-z]", limpio, re.I) and not re.match(r"^RA\d+[.\s]", limpio, re.I):
                    est["objetivos"].append(limpio)
                continue

        # Intercept "Enunciado:" for active blocks before debe_elim eats it
        if blk and (line == "Enunciado" or line.startswith("Enunciado:")):
            blk.pop("_skip_obj", None)
            blk.setdefault("enunciado_linea", len(blk.get("lineas", [])))
            continue

        if debe_elim(line):
            continue

        # Section headers (RE_SEC3 > RE_SEC2 > RE_SEC1)
        m3 = RE_SEC3.match(line)
        m2 = RE_SEC2.match(line)
        m1 = RE_SEC1.match(line)

        if m3 and current_sub:
            flush_blk()
            num = f"{current_sub['num']}.{m3.group(3)}"
            current_sub2 = {"num": num, "titulo": limpiar_titulo(m3.group(4)), "bloques": [], "subsecciones": []}
            current_sub.setdefault("subsecciones", []).append(current_sub2)
            continue

        if m2 and current_sec:
            flush_blk()
            num = f"{current_sec['num']}.{m2.group(2)}"
            current_sub = {"num": num, "titulo": limpiar_titulo(m2.group(3)), "bloques": [], "subsecciones": []}
            current_sub2 = None
            current_sec.setdefault("subsecciones", []).append(current_sub)
            continue

        if m1:
            flush_blk()
            current_sec = {"num": m1.group(1), "titulo": limpiar_titulo(m1.group(2)), "bloques": [], "subsecciones": []}
            current_sub = None
            current_sub2 = None
            est["secciones"].append(current_sec)
            continue

        if line == "Introducción":
            flush_blk()
            current_sec = {"num": "", "titulo": "Introducción", "bloques": [], "subsecciones": []}
            current_sub = None
            current_sub2 = None
            est["secciones"].append(current_sec)
            continue

        if not current_sec:
            continue

        # Interaction references
        mi = RE_INTER.match(line)
        if mi:
            flush_blk()
            n = int(mi.group(1))
            label = (mi.group(2) or "").strip()
            inter = interacciones.get(n, {})
            if inter.get("tipo") == "opciones":
                _add_inter_opciones(n, label, inter)
            else:
                for b in expandir_interaccion(n, interacciones, label):
                    activos().append(b)
            continue

        # Standalone block markers (e.g. "Para saber más", "Nota")
        if line in BLOQUES_ESP:
            flush_blk()
            tipo = _tipo_map.get(line, "ejemplo")
            etiqueta = normalizar_etiqueta_actividad(line) if RE_ACTIVITY_LABEL.match(line) else line
            blk = {"tipo": tipo, "etiqueta": etiqueta, "lineas": [], "_estilo": "text"}
            continue

        # Activity labels with optional number: "Actividad colaborativa 1", "Aplicación práctica 1"
        if RE_ACTIVITY_LABEL.match(line):
            flush_blk()
            m_act = RE_ACTIVITY_LABEL.match(line)
            base = m_act.group(1).strip()
            tipo = _tipo_map.get(base, "tarea")
            etiqueta = normalizar_etiqueta_actividad(line)
            blk = {"tipo": tipo, "etiqueta": etiqueta, "lineas": [], "_estilo": "text"}
            continue

        # URL / image blocks
        if RE_URL.match(line):
            if _blk_info_activo_pdf():
                blk.setdefault("lineas", []).append(rich_obj(line, [{"text": line, "link": line}]))
            else:
                flush_blk()
                activos().append({"tipo": "url_imagen" if _es_url_imagen(line) else "url", "url": line})
            continue

        if line.startswith("Pie de imagen:"):
            if _blk_info_activo_pdf():
                blk.setdefault("lineas", []).append(rich_obj(line, [{"text": line}]))
            else:
                activos().append({"tipo": "pie_imagen", "texto": line[len("Pie de imagen:"):].strip()})
            continue

        if re.match(r"^Descripci[oó]n de (la )?imagen:", line):
            if _blk_info_activo_pdf():
                blk.setdefault("lineas", []).append(rich_obj(line, [{"text": line}]))
            else:
                activos().append({
                    "tipo": "desc_imagen",
                    "texto": re.sub(r"^Descripci[oó]n de (la )?imagen:\s*", "", line),
                })
            continue

        # Content within active block
        if blk:
            if line.startswith(("Objetivos:", "Objetivo:")):
                blk["_skip_obj"] = True
            elif not blk.get("_skip_obj") and not _es_marcador_solucion(line) and not line.startswith("Duración:"):
                blk.setdefault("lineas", []).append(rich_obj(line, [{"text": line}]))
            continue

        # Regular paragraph
        activos().append({"tipo": "parrafo", "texto": rich_obj(line, [{"text": line}])})

    flush_blk()
    return est


# =============================================================================
# Parser DOCX fuente
# =============================================================================

def parsear_docx_fuente(docx_path: Path, interacciones: dict[int, dict]) -> dict:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph as DParagraph
    from docx.table import Table as DTable

    doc = Document(str(docx_path))

    relmap = {}
    try:
        for rel in doc.part.rels.values():
            if str(rel.reltype).endswith('/hyperlink'):
                relmap[rel.rId] = rel.target_ref
    except Exception:
        relmap = {}

    def _runs_desde_para(para) -> list[dict]:
        runs: list[dict] = []

        def add_run_xml(r_el, link_target: str | None = None):
            texts = []
            for t_el in r_el.findall('.//' + qn('w:t')):
                texts.append(t_el.text or '')
            if not texts:
                return
            text = ''.join(texts)
            rpr = r_el.find(qn('w:rPr'))
            bold_r = bool(rpr is not None and rpr.find(qn('w:b')) is not None)
            italic_r = bool(rpr is not None and rpr.find(qn('w:i')) is not None)
            color_val = ''
            if rpr is not None:
                color_el = rpr.find(qn('w:color'))
                if color_el is not None:
                    color_val = color_el.get(qn('w:val')) or ''
            item = {'text': text}
            if bold_r:
                item['bold'] = True
            if italic_r:
                item['italic'] = True
            if color_val and color_val.lower() not in {'auto', '000000'}:
                item['color'] = color_val
            if link_target:
                item['link'] = link_target
            runs.append(item)

        for child_el in para._p.iterchildren():
            if child_el.tag == qn('w:r'):
                add_run_xml(child_el)
            elif child_el.tag == qn('w:hyperlink'):
                rid = child_el.get(qn('r:id'))
                target = relmap.get(rid, '')
                for r_el in child_el.findall(qn('w:r')):
                    add_run_xml(r_el, target)

        if not runs and para.text:
            runs = [{'text': para.text}]
        return _trim_runs(runs)

    def _rich_para(para) -> dict:
        # v13: conservamos el texto ancla en el párrafo y guardamos las URLs
        # para volcarlas como líneas aparte al final del párrafo. Antes se
        # sustituía el ancla por la URL y el resultado quedaba ilegible.
        runs = _runs_desde_para(para)
        texto = _norm_line(''.join(r.get('text', '') for r in runs) if runs else para.text)
        obj = rich_obj(texto, runs or [{'text': texto}])
        links = []
        for r in runs or []:
            link = str(r.get('link', '')).strip()
            if link and RE_URL.match(link) and link not in links:
                links.append(link)
        if links:
            obj['links'] = links
        return obj

    est = {
        "titulo_unidad": "",
        "titulo_modulo": "",
        "objetivos": [],
        "secciones": [],
    }

    current_sec = None
    current_sub = None
    current_sub2 = None
    sec_count = 0
    sub_count = 0
    sub2_count = 0
    en_objetivos = False
    ignorar_hasta_contenido = False
    blk = None

    def activos() -> list:
        if current_sub2:
            return current_sub2["bloques"]
        if current_sub:
            return current_sub["bloques"]
        if current_sec:
            return current_sec["bloques"]
        return []

    def cerrar_ejercicio_pendiente():
        nonlocal blk

        if not blk or blk.get("tipo") != "_ejercicio_pendiente":
            return False

        inter = blk.get("inter", {})
        letras = "abcdefgh"
        opciones = []

        opciones_rich = inter.get("opciones_rich", [])
        if opciones_rich:
            for j, opt_rich in enumerate(opciones_rich):
                if j < len(letras):
                    opciones.append({"letra": letras[j], "texto": opt_rich})
        else:
            for j, opt in enumerate(inter.get("opciones", [])):
                opciones.append({
                    "letra": letras[j],
                    "texto": re.sub(r"^[a-h]\)\s*", "", opt.strip()),
                })

        tipo_final = blk.get("_tipo_final", "tarea")
        enunciado_text = inter.get("enunciado", "").strip()

        # Some interaction cells begin with a generic UI label ("Opciones de respuesta:",
        # "Opciones:", etc.) before the lettered choices. That label is not the actual
        # question enunciado, so we must not let it override the real body paragraphs.
        _RE_GENERIC_INTER_LABEL = re.compile(
            r"^(Opciones?(?:\s+de\s+respuesta)?|Enunciado|Instrucci[oó]n|Desplegables?)\s*:?\s*$",
            re.I,
        )
        _inter_enunciado_es_etiqueta = bool(
            not enunciado_text or _RE_GENERIC_INTER_LABEL.match(enunciado_text)
        )

        lineas_cuerpo_raw = [x for x in blk.get("lineas", []) if rich_text(x).strip()]

        if enunciado_text and not _inter_enunciado_es_etiqueta and not lineas_cuerpo_raw:
            # The interaction cell has a genuine question text AND there are no body
            # paragraphs — use the cell content as the enunciado.
            enunciado_lineas = [rich_obj(enunciado_text, [{"text": enunciado_text}])]
            lineas_cuerpo = []
        else:
            # Body paragraphs take priority: they contain the real question/exercise text.
            # The interaction cell may only have a generic label ("Opciones de respuesta:")
            # that is not the actual enunciado.
            enunciado_lineas = []
            lineas_cuerpo = lineas_cuerpo_raw

        entry = {
            "tipo": tipo_final,
            "etiqueta": blk.get("etiqueta", "Aplicación práctica"),
            "lineas": enunciado_lineas + lineas_cuerpo,
            "opciones": opciones,
        }
        if tipo_final == "aplicacion_practica":
            entry["solucion"] = inter.get("solucion_rich") or inter.get("solucion", "")
            entry["feedback"] = inter.get("feedback_rich") or inter.get("feedback", "")
        activos().append(entry)

        blk = None
        return True

    def flush():
        nonlocal blk

        if not blk:
            return

        if blk.get("tipo") == "_ejercicio_pendiente":
            cerrar_ejercicio_pendiente()
            return

        if blk.get("lineas") or blk.get("tipo") in {
            "hilo_conductor", "video", "ejemplo", "importante",
            "definicion", "sabias_que",
        }:
            b = {k: v for k, v in blk.items() if not k.startswith("_")}
            activos().append(b)

        blk = None

    def add(b: dict):
        flush()
        activos().append(b)

    def nueva_sec(num: str, titulo: str):
        nonlocal current_sec, current_sub, current_sub2, sec_count, sub_count, sub2_count

        flush()
        if num:
            try:
                sec_count = max(sec_count, int(str(num)))
            except Exception:
                pass
        else:
            sec_count += 1
        sub_count = 0
        sub2_count = 0
        current_sec = {
            "num": str(sec_count),
            "titulo": limpiar_titulo(titulo),
            "bloques": [],
            "subsecciones": [],
        }
        est["secciones"].append(current_sec)
        current_sub = None
        current_sub2 = None

    def nueva_sub(num: str, titulo: str):
        nonlocal current_sub, current_sub2, sub_count, sub2_count

        flush()

        if not current_sec:
            nueva_sec("", "Introducción")

        if num:
            try:
                sub_count = max(sub_count, int(str(num).split(".")[-1]))
            except Exception:
                pass
        else:
            sub_count += 1
            num = f'{current_sec.get("num", "1")}.{sub_count}'
        sub2_count = 0

        current_sub = {
            "num": num,
            "titulo": limpiar_titulo(titulo),
            "bloques": [],
            "subsecciones": [],
        }
        current_sec["subsecciones"].append(current_sub)
        current_sub2 = None

    def nueva_sub2(num: str, titulo: str):
        nonlocal current_sub2, sub2_count

        flush()

        if not current_sub:
            return

        if num:
            try:
                sub2_count = max(sub2_count, int(str(num).split(".")[-1]))
            except Exception:
                pass
        else:
            sub2_count += 1
            num = f'{current_sub.get("num", "1.1")}.{sub2_count}'

        current_sub2 = {
            "num": num,
            "titulo": limpiar_titulo(titulo),
            "bloques": [],
            "subsecciones": [],
        }
        current_sub["subsecciones"].append(current_sub2)

    special_styles = {
        "Hilo conductor": "hilo_conductor",
        "Ejemplo": "ejemplo",
        "Vídeo": "video",
        "Actividad colaborativa": "actividad_complementaria",
        "Importante": "importante",
        "Nota": "nota",
        "Consejo": "consejo",
        "Para saber más": "para_saber_mas",
        "Aplicación práctica": "aplicacion_practica",
        "Actividad de aprendizaje": "aplicacion_practica",
        "Sabiasque": "sabias_que",
        "Definición": "definicion",
    }

    list_styles = {
        "List Paragraph",
        "Viñeta nvl1 1d",
        "Viñeta nvl2 1d",
        "Vietanvl11d",
        "Vietanvl21d",
        "Ejemplos - Viñeta nvl1",
        "Ejemplos-Vietanvl1",
        "Recuerda - Viñeta nvl1",
        "Recuerda-Vietanvl1",
    }

    def _aplanar_elementos(body):
        elementos = []
        for child in body:
            if child.tag == qn("w:p"):
                elementos.append(child)
            elif child.tag == qn("w:tbl"):
                elementos.append(child)
        return elementos

    def _load_num_fmt_map() -> dict:
        """Returns {(numId_str, ilvl_str): numFmt_str} from the document's numbering part."""
        result: dict = {}
        try:
            nbp = doc.part.numbering_part
            if nbp is None:
                return result
            nxml = nbp._element
            abstract_defs: dict = {}
            for abn in nxml.findall(qn("w:abstractNum")):
                abs_id = abn.get(qn("w:abstractNumId"))
                if abs_id is None:
                    continue
                levels: dict = {}
                for lvl in abn.findall(qn("w:lvl")):
                    ilvl = lvl.get(qn("w:ilvl"))
                    fmt_el = lvl.find(qn("w:numFmt"))
                    if fmt_el is not None:
                        levels[ilvl] = fmt_el.get(qn("w:val"), "")
                abstract_defs[abs_id] = levels
            for num in nxml.findall(qn("w:num")):
                num_id = num.get(qn("w:numId"))
                if num_id is None:
                    continue
                abs_ref = num.find(qn("w:abstractNumId"))
                if abs_ref is not None:
                    abs_id = abs_ref.get(qn("w:val"))
                    if abs_id in abstract_defs:
                        for ilvl, fmt in abstract_defs[abs_id].items():
                            result[(num_id, ilvl)] = fmt
        except Exception:
            pass
        return result

    _num_fmt_map = _load_num_fmt_map()
    _numpr_idx: dict = {}  # (numId, ilvl) → next index (0-based)

    def _numpr_prefix(para_el) -> str:
        """Return the auto-number prefix for a paragraph with numPr, or '' if none."""
        try:
            pPr = para_el.find(qn("w:pPr"))
            if pPr is None:
                return ""
            numPr = pPr.find(qn("w:numPr"))
            if numPr is None:
                return ""
            num_id_el = numPr.find(qn("w:numId"))
            ilvl_el = numPr.find(qn("w:ilvl"))
            if num_id_el is None or ilvl_el is None:
                return ""
            num_id = num_id_el.get(qn("w:val"), "0")
            ilvl = ilvl_el.get(qn("w:val"), "0")
            if num_id == "0":
                return ""
            key = (num_id, ilvl)
            fmt = _num_fmt_map.get(key, "")
            idx = _numpr_idx.get(key, 0)
            _numpr_idx[key] = idx + 1
            if fmt == "lowerLetter":
                return chr(ord('a') + idx % 26) + ") "
            if fmt == "upperLetter":
                return chr(ord('A') + idx % 26) + ") "
            if fmt in ("decimal", "decimalZero"):
                return str(idx + 1) + ". "
            if fmt == "lowerRoman":
                _roman = ["i","ii","iii","iv","v","vi","vii","viii","ix","x"]
                return (_roman[idx] if idx < len(_roman) else str(idx+1)) + ". "
            if fmt == "upperRoman":
                _roman = ["I","II","III","IV","V","VI","VII","VIII","IX","X"]
                return (_roman[idx] if idx < len(_roman) else str(idx+1)) + ". "
        except Exception:
            pass
        return ""

    def _es_inicio_bloque_cualquiera(t, s):
        if s in special_styles:
            return True
        if t in BLOQUES_ESP:
            return True
        if RE_ACTIVITY_LABEL.match(t):
            return True
        base_esp = {"Nota", "Ejemplo", "Vídeo", "Video", "Importante", "Recuerda", "Sabías que"}
        for b in base_esp:
            if re.match(rf"^{b}(?:\s+\d+)?$", t, re.I):
                return True
        return False

    # Tipos de bloque informativo (texto) que deben absorber párrafos siguientes
    _TIPOS_BLOQUE_INFO = {
        "nota", "ejemplo", "para_saber_mas", "video", "importante",
        "sabias_que", "consejo", "definicion", "hilo_conductor", "recuerda",
    }

    def _blk_texto_activo():
        return (blk is not None
                and blk.get("_estilo") == "_texto"
                and blk.get("tipo") in _TIPOS_BLOQUE_INFO)

    for child in _aplanar_elementos(doc.element.body):
        if child.tag == qn("w:tbl"):
            tbl = DTable(child, doc)
            filas = []
            for row in tbl.rows:
                fila = []
                seen_tc = set()
                for cell in row.cells:
                    cid = id(cell._tc)
                    if cid in seen_tc:
                        continue
                    seen_tc.add(cid)
                    cell_lines = []
                    for para in cell.paragraphs:
                        ptxt = _norm_line(para.text)
                        if ptxt:
                            cell_lines.append(ptxt)
                    fila.append("\n".join(cell_lines))
                filas.append(fila)
            if any(any(c for c in row) for row in filas):
                tabla_blk = {"tipo": "tabla", "filas": filas}
                _tipos_absorb = {
                    "nota", "ejemplo", "para_saber_mas", "video", "importante",
                    "sabias_que", "consejo", "definicion", "hilo_conductor", "recuerda",
                    "actividad_complementaria", "tarea", "aplicacion_practica",
                }
                if _blk_texto_activo() or (blk is not None and blk.get("tipo") in _tipos_absorb):
                    blk.setdefault("lineas", []).append(tabla_blk)
                else:
                    flush()
                    activos().append(tabla_blk)
            continue
        is_para = child.tag == qn("w:p")
        if not is_para:
            continue

        para = DParagraph(child, doc)
        style = _style_name(para)
        txt = _norm_line(para.text)
        rich = _rich_para(para)
        bold = _is_bold(para)

        # Prepend auto-number prefix (a), b), 1., etc.) when the paragraph uses numPr
        # and the actual number character is NOT in the paragraph's <w:t> content.
        _np = _numpr_prefix(child)
        if _np and txt and not txt.startswith(_np.strip()):
            txt = _np + txt
            runs = [{"text": _np}] + list(rich_runs(rich))
            rich = rich_obj(txt, runs)

        if es_solo_pua(txt):
            continue


        if txt in {"", "Cambio de pantalla", "Específicos"} and style not in special_styles:
            # Si venimos recogiendo un bloque especial (tarea, ejemplo, etc.),
            # un cambio de pantalla marca el cierre del bloque. Antes se hacía
            # continue directo y algunas tareas justo antes del resumen se perdían.
            if txt == "Cambio de pantalla" and blk:
                flush()
            continue

        if _es_cabecera_objetivos(txt) and style not in special_styles and not current_sec:
            en_objetivos = True
            ignorar_hasta_contenido = False
            nueva_intro = _tipo_objetivos_intro(txt)
            if "general" in txt.lower():
                est["_current_obj_type"] = "general"
            else:
                est["_current_obj_type"] = "especifico"
            if "específico" in txt.lower() or "especifico" in txt.lower():
                est["_objectives_intro"] = nueva_intro
            continue

        if _es_cabecera_no_contenido(txt) and not current_sec:
            en_objetivos = False
            ignorar_hasta_contenido = True
            continue

        if blk and blk.get("_estilo") in {"Aplicación práctica", "_texto"} and (txt == "Enunciado" or txt.startswith("Enunciado:")):
            blk.pop("_skip_obj", None)
            blk.setdefault("enunciado_linea", len(blk.get("lineas", [])))
            # Extract text after "Enunciado:" label — the paragraph often contains
            # both the label and the enunciado text in the same run.
            if txt.startswith("Enunciado:"):
                content = txt[len("Enunciado:"):].strip()
                if content:
                    blk.setdefault("lineas", []).append(rich_obj(content, [{"text": content}]))
            continue

        if debe_elim(txt):
            # Inside an open "Actividad colaborativa" block, preserve all content
            # so the full activity description is transferred to paper format.
            if not (blk is not None and blk.get("_estilo") == "Actividad colaborativa"):
                continue

        if blk and blk.get("tipo") == "_ejercicio_pendiente":
            es_fin = False

            if style.startswith("Heading"):
                es_fin = True
            elif txt == "Enunciado" or txt.startswith("Enunciado:"):
                # Enunciado content is always part of the current exercise block,
                # regardless of its paragraph style (which may be "Aplicación práctica").
                pass  # es_fin stays False
            elif style in special_styles:
                es_fin = True
            elif RE_INTER.match(txt):
                es_fin = True
            elif txt in BLOQUES_ESP:
                es_fin = True
            elif txt.startswith(("Duración:", "Objetivo:")):
                continue

            if es_fin:
                cerrar_ejercicio_pendiente()
            else:
                item = rich
                if txt == "Enunciado":
                    continue
                if txt.startswith("Enunciado:"):
                    stripped = txt[len("Enunciado:"):].strip()
                    item = rich_obj(stripped, [{"text": stripped}])
                if txt and not re.match(r"^(Soluci[oó]n|Feedback|Retroalimentaci[oó]n):", txt, re.I):
                    blk["lineas"].append(item)
                continue

        mtitulo = _match_titulo_unidad(txt)
        if mtitulo and (style in {"Title", "Ttulo"} or style.startswith("_TITULO UNIDAD") or not est["titulo_unidad"]):
            etiqueta_unidad, n_unidad, titulo_modulo = mtitulo
            est["titulo_unidad"] = f"{etiqueta_unidad} {n_unidad}"
            if titulo_modulo and not est["titulo_modulo"]:
                est["titulo_modulo"] = titulo_modulo
            continue

        if style in {"Title", "Ttulo"} or style.startswith("_TITULO UNIDAD"):
            if est["titulo_unidad"] and not est["titulo_modulo"] and txt:
                est["titulo_modulo"] = limpiar_titulo(txt)
            continue

        if est["titulo_unidad"] and not est["titulo_modulo"] and txt and not _es_cabecera_objetivos(txt):
            if style in {"Title", "_TITULO UNIDAD 2"} or (not RE_SEC1.match(txt) and not RE_SEC2.match(txt) and len(txt) < 100 and not RE_URL.match(txt)):
                est["titulo_modulo"] = limpiar_titulo(txt)
                continue

        if en_objetivos:
            if (style.startswith("Heading") or RE_SEC1.match(txt) or txt == "Introducción" 
                or _es_cabecera_no_contenido(txt) or _es_inicio_bloque_cualquiera(txt, style)):
                en_objetivos = False
                if _es_cabecera_no_contenido(txt):
                    ignorar_hasta_contenido = True
                    continue
            elif txt and not txt.startswith("CE ") and not re.match(r"^[a-h]\)\s*Se han", txt, re.I):
                if not re.match(r"^RA\d+[.\s]", txt, re.I):
                    if est.get("_current_obj_type") == "general":
                        est["objetivo_general"] = rich
                    else:
                        est.setdefault("objetivos", []).append(_limpiar_vineta_rich(rich))
                continue

        if ignorar_hasta_contenido and not current_sec:
            if style.startswith("Heading") or txt == "Introducción" or RE_SEC1.match(txt):
                ignorar_hasta_contenido = False
            else:
                continue

        m3h = RE_SEC3.match(txt)
        m2h = RE_SEC2.match(txt)
        m1h = RE_SEC1.match(txt)

        if style == "Heading 1" or style == "1 Título nvl1":
            if txt == "Introducción":
                nueva_sec("", "Introducción")
                continue
            if m3h and current_sub:
                nueva_sub2(f"{current_sub.get('num', current_sec.get('num', m3h.group(1)) + '.' + m3h.group(2))}.{m3h.group(3)}", m3h.group(4))
                continue
            if m2h and current_sec:
                nueva_sub(f"{current_sec.get('num', m2h.group(1))}.{m2h.group(2)}", m2h.group(3))
                continue
            if m1h:
                nueva_sec(m1h.group(1), m1h.group(2))
            else:
                nueva_sec("", txt)
            continue

        if style == "Heading 2" or style == "2 Título nvl2":
            if m3h and current_sub:
                nueva_sub2(f"{current_sub.get('num', current_sec.get('num', m3h.group(1)) + '.' + m3h.group(2))}.{m3h.group(3)}", m3h.group(4))
            elif m2h:
                nueva_sub(f"{current_sec.get('num', m2h.group(1))}.{m2h.group(2)}", m2h.group(3))
            else:
                nueva_sub("", txt)
            continue

        if style == "Heading 3" or style == "3 Título nvl3":
            if m3h:
                nueva_sub2(f"{current_sub.get('num', current_sec.get('num', m3h.group(1)) + '.' + m3h.group(2))}.{m3h.group(3)}", m3h.group(4))
            elif m2h:
                nueva_sub(f"{current_sec.get('num', m2h.group(1))}.{m2h.group(2)}", m2h.group(3))
            elif m1h:
                nueva_sec(m1h.group(1), m1h.group(2))
            else:
                nueva_sub2("", txt)
            continue

        if current_sec is None and txt == "Introducción":
            nueva_sec("", "Introducción")
            continue

        # Dentro de una tarea de evaluación, las líneas numeradas del enunciado
        # (1., 2., 3...) son instrucciones, no títulos de nuevas secciones.
        if blk and blk.get("_estilo") in {"Aplicación práctica", "_texto"} and txt and not RE_INTER.match(txt):
            # Seguimos dentro del mismo recurso práctico mientras el estilo continúe
            # siendo "Aplicación práctica" o estemos en un bloque de texto aplanado (_texto).
            # Cerramos el bloque solo si:
            #   - Es un Heading (nueva sección)
            #   - Es un special_style que además es la CABECERA de ese bloque (txt == label del estilo),
            #     lo que indica que empieza un bloque especial nuevo, no líneas de contenido del mismo bloque.
            # NOTE: RE_INTER is excluded here so blk remains intact for the rescue code
            # at the RE_INTER handler below, which needs blk.lineas to carry the enunciado.
            es_nueva_cabecera_especial = (
                style in special_styles
                and style != blk.get("_estilo")
                and _es_inicio_bloque_cualquiera(txt, style)
            )
            if style.startswith("Heading") or es_nueva_cabecera_especial:
                flush()
            else:
                if _es_marcador_solucion(txt):
                    if _bloque_admite_solucion(blk):
                        blk.setdefault("solucion_lineas", [])
                        # Conservamos la etiqueta de solución como parte del bloque
                        # solo cuando el recurso es Aplicación práctica.
                        blk["_capturando_solucion"] = True
                        blk["solucion_lineas"].append(rich)
                    else:
                        blk["_skip_solucion"] = True
                    continue
                if blk.get("_skip_solucion") and not _bloque_admite_solucion(blk):
                    continue
                if blk.get("_capturando_solucion") and _bloque_admite_solucion(blk):
                    blk.setdefault("solucion_lineas", []).append(rich)
                    continue
                if txt == "Enunciado" or txt.startswith("Enunciado:"):
                    blk.pop("_skip_obj", None)
                    blk.setdefault("enunciado_linea", len(blk.get("lineas", [])))
                elif txt.startswith(("Objetivos:", "Objetivo:")):
                    blk["_skip_obj"] = True
                elif not blk.get("_skip_obj") and txt != "Enunciado" and not txt.startswith(("Duración:", "Objetivo:", "Objetivos:", "Enunciado:")) and not debe_elim(txt):
                    if style in list_styles:
                        nivel_v = 2 if "nvl2" in style.lower() or "21" in style else 1
                        blk.setdefault("lineas", []).append(_prefix_rich(f"{VINETA_SIM[nivel_v]} ", _limpiar_vineta_rich(rich)))
                    else:
                        blk.setdefault("lineas", []).append(rich)
                continue

        m3 = RE_SEC3.match(txt)
        if m3 and current_sec:
            nueva_sub2(f"{current_sub.get("num", current_sec.get("num", m3.group(1)) + "." + m3.group(2))}.{m3.group(3)}", m3.group(4))
            continue

        m2 = RE_SEC2.match(txt)
        if m2 and current_sec:
            nueva_sub(f"{current_sec.get("num", m2.group(1))}.{m2.group(2)}", m2.group(3))
            continue

        m1 = RE_SEC1.match(txt)
        if m1 and len(m1.group(2)) < 120:
            nueva_sec(m1.group(1), m1.group(2))
            continue

        if not current_sec:
            continue

        mi = RE_INTER.match(txt)
        if mi:
            n = int(mi.group(1))
            label = (mi.group(2) or "").strip()
            inter = interacciones.get(n, {})

            # When the preceding block ("Aplicación práctica" / "tarea") has
            # accumulated the enunciado, we must NOT flush it away before building
            # the exercise entry — instead, rescue its lineas and discard the blk.
            lineas_rescatadas: list = []
            etiqueta_rescatada: str = ""
            if (inter.get("tipo") == "opciones"
                    and blk
                    and blk.get("_estilo") in {"Aplicación práctica", "tarea"}
                    and not blk.get("_ejercicio_pendiente")):
                lineas_rescatadas = [x for x in blk.get("lineas", []) if rich_text(x).strip()]
                etiqueta_rescatada = blk.get("etiqueta", "")
                blk = None  # discard without adding to activos

            flush()  # flush any remaining unrelated block

            if inter.get("tipo") == "opciones":
                mt_eval = re.search(r"Actividad de evaluaci[oó]n", label, re.I)
                mt_apr = re.search(r"Actividad de aprendizaje\s+(\d+)", label, re.I)
                mt_aprac = re.search(r"Aplicaci[oó]n pr[aá]ctica\s*(\d*)", label, re.I)

                if mt_eval:
                    # "Actividad de evaluación X CE y" (inline) → "Actividad X" sin solución
                    _tipo_final = "tarea"
                    num_eval = re.search(r"(\d+)", label[mt_eval.end():]) or re.search(r"(\d+)", label)
                    _etiqueta = f"Actividad {num_eval.group(1)}" if num_eval else f"Actividad {n}"
                elif mt_apr:
                    # "Actividad de aprendizaje X" (test) → "Aplicación práctica X" CON solución
                    _tipo_final = "aplicacion_practica"
                    num_apr = mt_apr.group(1).strip()
                    _etiqueta = f"Aplicación práctica {num_apr}".strip() if num_apr else "Aplicación práctica"
                elif mt_aprac:
                    _tipo_final = "tarea"
                    num_apr = mt_aprac.group(1).strip()
                    _etiqueta = f"Tarea {num_apr}".strip() if num_apr else "Tarea"
                else:
                    _tipo_final = "tarea"
                    _etiqueta = label.strip() or etiqueta_rescatada or "Tarea"

                blk = {
                    "tipo": "_ejercicio_pendiente",
                    "_estilo": "_ejercicio",
                    "_tipo_final": _tipo_final,
                    "etiqueta": _etiqueta,
                    "lineas": lineas_rescatadas,  # enunciado from preceding block
                    "inter": inter,
                }
            else:
                for b in expandir_interaccion(n, interacciones, label):
                    activos().append(b)
            continue

        if RE_URL.match(txt):
            flush()
            activos().append({"tipo": "url_imagen" if _es_url_imagen(txt) else "url", "url": txt})
            continue

        if txt.startswith("Pie de imagen:"):
            flush()
            activos().append({"tipo": "pie_imagen", "texto": txt[len("Pie de imagen:"):].strip()})
            continue

        if re.match(r"^Descripci[oó]n de (la )?imagen:", txt, re.I):
            flush()
            activos().append({"tipo": "desc_imagen",
                              "texto": re.sub(r"^Descripci[oó]n de (la )?imagen:\s*", "", txt)})
            continue

        if style in special_styles:
            if style == "Sabiasque":
                if txt in {"Sabías que…", "Sabías que...", "Definición", "Importante"}:
                    flush()
                    tipo_real = {
                        "Sabías que…": "sabias_que",
                        "Sabías que...": "sabias_que",
                        "Definición": "definicion",
                        "Importante": "importante",
                    }.get(txt, "sabias_que")

                    blk = {
                        "tipo": tipo_real,
                        "etiqueta": txt,
                        "lineas": [],
                        "_estilo": style,
                    }
                else:
                    if not blk or blk.get("_estilo") != style:
                        flush()
                        blk = {
                            "tipo": "sabias_que",
                            "etiqueta": "Sabías que…",
                            "lineas": [],
                            "_estilo": style,
                        }
                    blk["lineas"].append(_limpiar_vineta_rich(rich))
                continue

            if style == "Definición":
                if txt == "Definición":
                    flush()
                    blk = {
                        "tipo": "definicion",
                        "etiqueta": "Definición",
                        "lineas": [],
                        "_estilo": style,
                    }
                elif blk and blk.get("_estilo") == style:
                    blk["lineas"].append(rich)
                else:
                    flush()
                    blk = {
                        "tipo": "definicion",
                        "etiqueta": "Definición",
                        "lineas": [rich],
                        "_estilo": style,
                    }
                continue

            if style == "Hilo conductor":
                if txt == "Hilo conductor":
                    flush()
                    blk = {
                        "tipo": "hilo_conductor",
                        "etiqueta": "Hilo conductor",
                        "lineas": [],
                        "_estilo": style,
                    }
                elif blk and blk.get("_estilo") == style:
                    blk["lineas"].append(rich)
                continue

            if style == "Ejemplo":
                if txt == "Ejemplo":
                    flush()
                    blk = {
                        "tipo": "ejemplo",
                        "etiqueta": "Ejemplo",
                        "lineas": [],
                        "_estilo": style,
                    }
                elif re.match(r"^Descripci[oó]n de (la )?imagen:", txt):
                    add({
                        "tipo": "desc_imagen",
                        "texto": re.sub(r"^Descripci[oó]n de (la )?imagen:\s*", "", txt),
                    })
                elif blk and blk.get("_estilo") == style:
                    blk["lineas"].append(rich)
                continue

            if style == "Vídeo":
                if txt == "Vídeo" and _links_from_rich(rich) and blk and blk.get("_estilo") == style:
                    # Es el botón/enlace del recurso, no un nuevo bloque de vídeo.
                    blk["lineas"].append(rich)
                elif txt == "Vídeo":
                    flush()
                    blk = {
                        "tipo": "video",
                        "etiqueta": "Vídeo",
                        "lineas": [],
                        "_estilo": style,
                    }
                elif blk and blk.get("_estilo") == style:
                    blk["lineas"].append(rich)
                continue

            if style == "Importante":
                if txt == "Importante" or not blk or blk.get("_estilo") != style:
                    flush()
                    blk = {
                        "tipo": "importante",
                        "etiqueta": "Importante",
                        "lineas": [],
                        "_estilo": style,
                    }
                    if txt and txt != "Importante":
                        blk["lineas"].append(rich)
                else:
                    blk["lineas"].append(rich)
                continue

            if style in {"Nota", "Consejo", "Para saber más"}:
                _tipo_label_map = {"Nota": "nota", "Consejo": "consejo", "Para saber más": "para_saber_mas"}
                tipo_real = _tipo_label_map[style]
                label = style
                if txt == label or not blk or blk.get("_estilo") != style:
                    flush()
                    blk = {
                        "tipo": tipo_real,
                        "etiqueta": label,
                        "lineas": [],
                        "_estilo": style,
                    }
                    if txt and txt != label:
                        blk["lineas"].append(rich)
                else:
                    blk["lineas"].append(rich)
                continue

            if style == "Actividad colaborativa":
                mnum = re.search(r"Actividad\s+colaborativa\s+(\d+)", txt, re.I)

                # En papel, las actividades complementarias NO deben incluir la
                # posible solución ni el feedback del material online. Cuando
                # aparece el marcador de solución, dejamos de recoger líneas
                # hasta que se cierre el bloque al detectar otro recurso/sección.
                if re.match(r"^(POSIBLE SOLUCI[OÓ]N|Soluci[oó]n|Retroalimentaci[oó]n|Feedback)\b", txt, re.I):
                    if blk and blk.get("_estilo") == style:
                        blk["_skip_solucion"] = True
                    continue

                if not blk or blk.get("_estilo") != style:
                    flush()
                    blk = {
                        "tipo": "actividad_complementaria",
                        "etiqueta": "Actividad complementaria",
                        "numero": mnum.group(1) if mnum else "",
                        "lineas": [],
                        "_estilo": style,
                        "_skip_solucion": False,
                    }
                elif mnum:
                    blk["numero"] = mnum.group(1)
                    continue
                elif blk.get("_skip_solucion"):
                    continue

                if txt and not re.match(r"^Actividad\s+colaborativa\b", txt, re.I):
                    blk["lineas"].append(rich)
                continue

            if style == "Aplicación práctica":
                if _es_marcador_solucion(txt):
                    if blk and blk.get("_estilo") == style:
                        if _bloque_admite_solucion(blk):
                            blk["_capturando_solucion"] = True
                            blk.setdefault("solucion_lineas", []).append(rich)
                        else:
                            blk["_skip_solucion"] = True
                    continue

                mt = re.search(r"(?:Tarea de evaluaci[oó]n|Aplicaci[oó]n pr[aá]ctica|Tarea)\s+(\d+)", txt, re.I)
                es_aplicacion = bool(re.match(r"^Aplicaci[oó]n pr[aá]ctica\b", txt, re.I))
                es_tarea = bool(re.match(r"^(?:Tarea|Tarea de evaluaci[oó]n)\b", txt, re.I))
                if not blk or blk.get("_estilo") != style:
                    flush()
                    num = mt.group(1) if mt else ""
                    if es_aplicacion:
                        blk = {
                            "tipo": "tarea",
                            "etiqueta": f"Tarea {num}".strip() if num else "Tarea",
                            "lineas": [],
                            "_estilo": style,
                            "_skip_solucion": False,
                        }
                    else:
                        blk = {
                            "tipo": "tarea",
                            "etiqueta": f"Tarea {num or '1'}",
                            "lineas": [],
                            "_estilo": style,
                            "_skip_solucion": False,
                        }
                elif mt and es_tarea:
                    blk["etiqueta"] = f"Tarea {mt.group(1)}"
                    blk["tipo"] = "tarea"
                    continue
                elif blk.get("_skip_solucion") and not _bloque_admite_solucion(blk):
                    continue
                elif blk.get("_capturando_solucion") and _bloque_admite_solucion(blk):
                    blk.setdefault("solucion_lineas", []).append(rich)
                    continue
                else:
                    if txt.startswith(("Objetivos:", "Objetivo:")):
                        blk["_skip_obj"] = True
                    elif txt == "Enunciado" or txt.startswith("Enunciado:"):
                        blk.pop("_skip_obj", None)
                        content = txt[len("Enunciado:"):].strip() if txt.startswith("Enunciado:") else ""
                        if content:
                            blk["lineas"].append(rich_obj(content, [{"text": content}]))
                    elif blk.get("_skip_obj"):
                        pass
                    elif txt and not re.match(r"^(Aplicaci[oó]n pr[aá]ctica|Tarea(?: de evaluaci[oó]n)?)\b", txt, re.I) and not txt.startswith("Duración:"):
                        if style in list_styles:
                            nivel_v = 2 if "nvl2" in style.lower() or "21" in style else 1
                            blk["lineas"].append(_prefix_rich(f"{VINETA_SIM[nivel_v]} ", _limpiar_vineta_rich(rich)))
                        else:
                            blk["lineas"].append(rich)
                continue

            if style == "Actividad de aprendizaje":
                if _es_marcador_solucion(txt):
                    if blk and blk.get("_estilo") == style:
                        if _bloque_admite_solucion(blk):
                            blk["_capturando_solucion"] = True
                            blk.setdefault("solucion_lineas", []).append(rich)
                        else:
                            blk["_skip_solucion"] = True
                    continue

                mt = re.search(r"(?:Actividad\s+de\s+aprendizaje|Aplicaci[oó]n\s+pr[aá]ctica)\s+(\d+)", txt, re.I)
                if not blk or blk.get("_estilo") != style:
                    flush()
                    num = mt.group(1) if mt else ""
                    blk = {
                        "tipo": "aplicacion_practica",
                        "etiqueta": f"Aplicación práctica {num}".strip() if num else "Aplicación práctica",
                        "lineas": [],
                        "_estilo": "Actividad de aprendizaje",
                        "_skip_solucion": False,
                    }
                elif blk.get("_skip_solucion") and not _bloque_admite_solucion(blk):
                    continue
                elif blk.get("_capturando_solucion") and _bloque_admite_solucion(blk):
                    blk.setdefault("solucion_lineas", []).append(rich)
                    continue
                else:
                    if txt.startswith(("Objetivos:", "Objetivo:")):
                        blk["_skip_obj"] = True
                    elif txt == "Enunciado" or txt.startswith("Enunciado:"):
                        blk.pop("_skip_obj", None)
                        content = txt[len("Enunciado:"):].strip() if txt.startswith("Enunciado:") else ""
                        if content:
                            blk["lineas"].append(rich_obj(content, [{"text": content}]))
                    elif blk.get("_skip_obj"):
                        pass
                    elif txt and not re.match(r"^Actividad\s+de\s+aprendizaje\b", txt, re.I) and not txt.startswith("Duración:"):
                        if style in list_styles:
                            nivel_v = 2 if "nvl2" in style.lower() or "21" in style else 1
                            blk["lineas"].append(_prefix_rich(f"{VINETA_SIM[nivel_v]} ", _limpiar_vineta_rich(rich)))
                        else:
                            blk["lineas"].append(rich)
                continue

        if blk and blk.get("_estilo") and style in list_styles:
            if txt:
                blk.setdefault("lineas", []).append(_prefix_rich("● ", _limpiar_vineta_rich(rich)))
            continue

        if blk and blk.get("_estilo") and blk.get("_estilo") != "_texto" and style not in special_styles:
            flush()

        def _obtener_base_bloque(t):
            if t in BLOQUES_ESP:
                return t
            if RE_ACTIVITY_LABEL.match(t):
                m = RE_ACTIVITY_LABEL.match(t)
                return m.group(1)
            base_esp = {"Nota", "Ejemplo", "Vídeo", "Video", "Importante", "Recuerda", "Sabías que"}
            for b in base_esp:
                if re.match(rf"^{b}(?:\s+\d+)?$", t, re.I):
                    if b in {"Vídeo", "Video"}:
                        return "Vídeo"
                    if b == "Sabías que":
                        return "Sabías que..."
                    return b
            return None

        base_bloque = _obtener_base_bloque(txt)
        if base_bloque:
            flush()

            tipo_map = {
                "Nota": "nota",
                "Ejemplo": "ejemplo",
                "Sabías que...": "sabias_que",
                "Sabías que…": "sabias_que",
                "Consejo": "consejo",
                "Definición": "definicion",
                "Hilo conductor": "hilo_conductor",
                "Para saber más": "para_saber_mas",
                "Vídeo": "video",
                "Video": "video",
                "Importante": "importante",
                "Recuerda": "recuerda",
                "Actividad complementaria": "actividad_complementaria",
                "Actividad colaborativa": "actividad_complementaria",
                "Actividad de evaluación": "tarea",
                "Aplicación práctica": "tarea",
                "Caso práctico": "tarea",
                "Ejercicio": "tarea",
                "Tarea": "tarea",
            }

            tipo_detectado = tipo_map.get(base_bloque, "ejemplo")
            etiqueta = normalizar_etiqueta_actividad(txt) if RE_ACTIVITY_LABEL.match(txt) else txt
            if base_bloque == "Aplicación práctica":
                etiqueta = re.sub(r"^Aplicaci[oó]n\s+pr[aá]ctica\b\s*", "Tarea ", etiqueta, flags=re.I).strip()
            blk = {
                "tipo": tipo_detectado,
                "etiqueta": etiqueta,
                "lineas": [],
                "_estilo": "_texto",
                "_skip_solucion": False,
            }
            continue

        if RE_URL.match(txt):
            if _blk_texto_activo():
                blk.setdefault("lineas", []).append(rich_obj(txt, [{"text": txt, "link": txt}]))
            else:
                add({
                    "tipo": "url_imagen" if _es_url_imagen(txt) else "url",
                    "url": txt,
                })
            continue

        if txt.startswith("Pie de imagen:"):
            if _blk_texto_activo():
                blk.setdefault("lineas", []).append(rich_obj(txt, [{"text": txt}]))
            else:
                add({"tipo": "pie_imagen", "texto": txt[len("Pie de imagen:"):].strip()})
            continue

        if re.match(r"^Descripci[oó]n de (la )?imagen:", txt):
            if _blk_texto_activo():
                blk.setdefault("lineas", []).append(rich_obj(txt, [{"text": txt}]))
            else:
                add({
                    "tipo": "desc_imagen",
                    "texto": re.sub(r"^Descripci[oó]n de (la )?imagen:\s*", "", txt),
                })
            continue

        if style in list_styles:
            nivel = 2 if "nvl2" in style.lower() or "21" in style else 1
            if _blk_texto_activo():
                blk.setdefault("lineas", []).append(
                    _prefix_rich(f"{VINETA_SIM[nivel]} ", _limpiar_vineta_rich(rich))
                )
            else:
                add({"tipo": "p_vineta", "texto": rich, "nivel": nivel})
            continue

        mo = RE_OPCION.match(txt)
        if mo:
            add({
                "tipo": "opcion_test_suelta",
                "letra": mo.group(1),
                "texto": mo.group(2),
            })
            continue

        if _blk_texto_activo():
            blk.setdefault("lineas", []).append(rich)
        elif blk and blk.get("tipo") == "actividad_complementaria" and not blk.get("_skip_solucion") and txt:
            blk["lineas"].append(rich)
        else:
            add({"tipo": "parrafo", "texto": rich})

    flush()

    if not est["titulo_unidad"]:
        est["titulo_unidad"] = "Unidad de aprendizaje 1"

    if not est["titulo_modulo"]:
        est["titulo_modulo"] = ""

    return est


# =============================================================================
# XML de bloques
# =============================================================================

def remove_empty_example_blocks(bloques: list[dict]) -> list[dict]:
    # 1. Agrupar URL/Imagen_XXX + Pie + Descripción
    out = []
    i = 0
    while i < len(bloques):
        b = bloques[i]
        
        es_recurso = False
        url_or_id = ""
        
        if b.get("tipo") in {"url", "url_imagen"}:
            es_recurso = True
            url_or_id = b.get("url", "")
        elif b.get("tipo") == "parrafo":
            texto = rich_text(b.get("texto", "")).strip()
            if re.match(r"^Imagen_\d+", texto, re.I):
                es_recurso = True
                url_or_id = texto
        
        if es_recurso:
            recurso = {
                "tipo": "recurso_agrupado",
                "url_or_id": url_or_id,
                "pie": None,
                "desc": None,
                "original_b": b,
            }
            i += 1
            while i < len(bloques):
                nxt = bloques[i]
                if nxt.get("tipo") == "pie_imagen":
                    recurso["pie"] = "Pie de imagen: " + nxt.get("texto", "").strip()
                    i += 1
                elif nxt.get("tipo") == "desc_imagen":
                    recurso["desc"] = "Descripción de la imagen: " + nxt.get("texto", "").strip()
                    i += 1
                else:
                    break
            out.append(recurso)
            continue
            
        out.append(b)
        i += 1
        
    # 2. Los recursos de imagen siempre van FUERA del bloque editorial anterior,
    #    justo después del cierre del recuadro. No se integran dentro.
    out2 = list(out)
            
    # 3. Eliminar Ejemplos vacios
    out3 = []
    for b in out2:
        if b.get("tipo") in {"ejemplo", "importante", "sabias_que", "nota", "consejo", "definicion", "hilo_conductor", "para_saber_mas", "video"}:
            tiene_contenido = False
            for line in b.get("lineas", []):
                if isinstance(line, dict) and line.get("tipo") == "recurso_agrupado_interno":
                    tiene_contenido = True
                    break
                txt = rich_text(line).strip()
                if txt:
                    tiene_contenido = True
                    break
            if not tiene_contenido:
                continue
        out3.append(b)
        
    return out3


def tabla_xml(filas: list) -> str:
    if not filas:
        return ""
    n_cols = max((len(row) for row in filas), default=1)
    if n_cols == 0:
        return ""
    # Ancho total del cuerpo de página A4 con márgenes estándar (≈15.5 cm)
    TABLE_W = 8748
    col_w = TABLE_W // n_cols
    grid = "".join(f'<w:gridCol w:w="{col_w}"/>' for _ in range(n_cols))
    # Padding de celda: 80 twips (≈ 1.4mm) por lado
    CELL_MAR = (
        '<w:tcMar>'
        '<w:top w:w="80" w:type="dxa"/>'
        '<w:left w:w="108" w:type="dxa"/>'
        '<w:bottom w:w="80" w:type="dxa"/>'
        '<w:right w:w="108" w:type="dxa"/>'
        '</w:tcMar>'
    )
    rows_xml = []
    for fila in filas:
        cells_xml = []
        for i in range(n_cols):
            raw = fila[i] if i < len(fila) else ""
            paras = raw.split("\n") if raw else [""]
            paras_xml = "".join(
                f'<w:p><w:pPr><w:pStyle w:val="Cuerpoparrafo"/></w:pPr>'
                f'<w:r><w:t xml:space="preserve">{esc(ln)}</w:t></w:r></w:p>'
                for ln in paras
            )
            cells_xml.append(
                f'<w:tc>'
                f'<w:tcPr><w:tcW w:w="{col_w}" w:type="dxa"/>{CELL_MAR}</w:tcPr>'
                f'{paras_xml}</w:tc>'
            )
        rows_xml.append("<w:tr>" + "".join(cells_xml) + "</w:tr>")
    return (
        '<w:tbl>'
        '<w:tblPr><w:tblStyle w:val="TableGrid"/>'
        f'<w:tblW w:w="{TABLE_W}" w:type="dxa"/>'
        '<w:tblBorders>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders></w:tblPr>'
        f'<w:tblGrid>{grid}</w:tblGrid>'
        + "".join(rows_xml)
        + '</w:tbl>'
    )


def bloques_xml(bloques: list[dict]) -> list[str]:
    out: list[str] = []

    for b in bloques:
        t = b.get("tipo", "")

        if t == "parrafo":
            texto = b.get("texto", "")
            texto_plano = rich_text(texto)
            if re.match(r"^Imagen_\d+", texto_plano, re.I):
                out.append(add_image_label_paragraph(texto_plano, "Normal"))
            elif RE_FORMULA.search(texto_plano) and len(texto_plano) < 200:
                out.append(p_formula(texto))
            else:
                out.append(p(texto, "Cuerpoparrafo"))
            _append_links_xml(out, texto)

        elif t == "p_vineta":
            texto = b.get("texto", "")
            out.append(p_vineta(texto, b.get("nivel", 1)))
            _append_links_xml(out, texto)

        elif t == "p_vineta_bold":
            out.append(p_vineta_bold(b.get("texto", ""), b.get("nivel", 1)))

        elif t == "p_vineta_ejemplo":
            out.append(p_vineta_ejemplo(b.get("texto", "")))

        elif t == "desplegable":
            titulo = b.get("titulo", "")
            desc = b.get("descripcion", "")
            if desc:
                out.append(p_desp(titulo, desc))
            else:
                out.append(p_vineta_bold(titulo))

        elif t == "desplegable_simple":
            out.append(p_desp(b.get("titulo_rich") or b.get("titulo", ""), b.get("contenido_rich") or b.get("contenido", "")))

        elif t == "desplegable_simple_n2":
            out.append(p_desp(b.get("titulo_rich") or b.get("titulo", ""), b.get("contenido_rich") or b.get("contenido", ""), 2))

        elif t == "desplegable_procedimental":
            out.append(p_vineta_bold(b.get("titulo_rich") or b.get("titulo", "")))
            body_r = b.get("body_rich") or b.get("body", "")
            if body_r:
                out.append(p_vineta(body_r, 2))
            subitems_rich = b.get("subitems_rich", [])
            for j, sub in enumerate(b.get("subitems", [])):
                sub_r = subitems_rich[j] if j < len(subitems_rich) else None
                out.append(p_vineta(sub_r if sub_r is not None else sub, 3 if body_r else 2))

        elif t == "recurso_agrupado":
            url_or_id = b.get("url_or_id", "")
            if re.match(r"^Imagen_\d+", url_or_id, re.I):
                out.append(add_image_label_paragraph(url_or_id, "Normal"))
            else:
                out.append(p_url_recurso(url_or_id))
            if b.get("pie"):
                out.append(add_image_label_paragraph(b["pie"], "Cuerpoparrafo"))
            if b.get("desc"):
                out.append(add_image_label_paragraph(b["desc"], "Normal"))

        elif t == "tabla":
            filas = b.get("filas", [])
            if filas:
                out.append(tabla_xml(filas))

        elif t == "url_imagen":
            out.append(p_url_imagen(b.get("texto") or b.get("url", "")))

        elif t == "url_recurso":
            out.append(p_url_recurso(b.get("texto") or b.get("url", "")))

        elif t == "pie_imagen":
            out.append(p_pie_imagen(b.get("texto", "")))

        elif t == "desc_imagen":
            out.append(p_desc_imagen(b.get("texto", "")))

        elif t == "desplegable_multi":
            out.append(p_vineta_bold(b.get("titulo", "").rstrip(":") + ":"))
            for item in b.get("items", []):
                out.append(p_vineta(item, 2))

        elif t in {
            "nota", "ejemplo", "sabias_que", "consejo", "definicion",
            "hilo_conductor", "para_saber_mas", "video", "importante",
        }:
            out.append(p(b.get("etiqueta", ""), "Recuerda-00lneainicio"))

            modo_lista = False

            for line in b.get("lineas", []):
                if isinstance(line, dict) and line.get("tipo") == "tabla":
                    out.append(tabla_xml(line.get("filas", [])))
                    continue
                if isinstance(line, dict) and line.get("tipo") == "recurso_agrupado_interno":
                    rec = line["recurso"]
                    url_or_id = rec.get("url_or_id", "")
                    if re.match(r"^Imagen_\d+", url_or_id, re.I):
                        out.append(add_image_label_paragraph(url_or_id, "Normal"))
                    else:
                        out.append(p(rich_obj(url_or_id, [{"text": url_or_id, "link": url_or_id}]), "Ejemplos-Cuerpoparrafo"))
                    if rec.get("pie"):
                        out.append(add_image_label_paragraph(rec["pie"], "Ejemplos-Cuerpoparrafo"))
                    if rec.get("desc"):
                        out.append(add_image_label_paragraph(rec["desc"], "Normal" if re.match(r"^Imagen_\d+", url_or_id, re.I) else "Ejemplos-Cuerpoparrafo"))
                    modo_lista = False
                    continue

                line_txt = rich_text(line).strip()

                if not line_txt:
                    continue

                if RE_URL.match(line_txt):
                    out.append(p_url_recurso(line_txt))
                    modo_lista = False
                    continue

                # Imágenes internas (Imagen_XX): dos párrafos Normal según PDF regla 8.
                # No resetean modo_lista: las imágenes no rompen la lista (PDF nota 8).
                if re.match(r"^Imagen_\d+", line_txt, re.I):
                    out.append(add_image_label_paragraph(line_txt, "Normal"))
                    continue

                if re.match(r"^Descripci[oó]n de (la )?imagen:", line_txt, re.I):
                    out.append(add_image_label_paragraph(line_txt, "Normal"))
                    continue

                # En recursos tipo Vídeo el online puede traer un botón con texto
                # visible "Vídeo" y el hipervínculo real en el run. En papel debe
                # quedar la URL, no otra viñeta con la palabra Vídeo.
                if _links_from_rich(line) and line_txt.lower() in {"vídeo", "video", "enlace", "ver vídeo", "ver video"}:
                    _append_links_xml(out, line)
                    modo_lista = False
                    continue

                if _parece_item_lista_en_bloque(line_txt, modo_lista):
                    out.append(p_vineta_ejemplo(line if isinstance(line, dict) else _limpiar_vineta_literal(line_txt)))
                    _append_links_xml(out, line)
                    continue

                out.append(p(line, "Ejemplos-Cuerpoparrafo"))
                _append_links_xml(out, line)
                modo_lista = _abre_modo_lista(line_txt)

            out.append(p("", "Recuerda-01lneafin"))

        elif t == "recuerda":
            out.append(p(b.get("etiqueta", "Recuerda"), "Recuerda-00lneainicio"))

            modo_lista_r = False
            for line in b.get("lineas", []):
                if isinstance(line, dict) and line.get("tipo") == "tabla":
                    out.append(tabla_xml(line.get("filas", [])))
                    continue
                if isinstance(line, dict) and line.get("tipo") == "recurso_agrupado_interno":
                    rec = line["recurso"]
                    url_or_id = rec.get("url_or_id", "")
                    if re.match(r"^Imagen_\d+", url_or_id, re.I):
                        out.append(add_image_label_paragraph(url_or_id, "Normal"))
                    else:
                        out.append(p(rich_obj(url_or_id, [{"text": url_or_id, "link": url_or_id}]), "Recuerda-Cuerpoparrafo"))
                    if rec.get("pie"):
                        out.append(add_image_label_paragraph(rec["pie"], "Recuerda-Cuerpoparrafo"))
                    if rec.get("desc"):
                        out.append(add_image_label_paragraph(rec["desc"], "Normal" if re.match(r"^Imagen_\d+", url_or_id, re.I) else "Recuerda-Cuerpoparrafo"))
                    modo_lista_r = False
                    continue

                line_txt = rich_text(line).strip()
                if not line_txt:
                    continue
                if RE_URL.match(line_txt):
                    out.append(p_url_recurso(line_txt))
                    modo_lista_r = False
                    continue
                if re.match(r"^Imagen_\d+", line_txt, re.I):
                    out.append(add_image_label_paragraph(line_txt, "Normal"))
                    continue
                if re.match(r"^Descripci[oó]n de (la )?imagen:", line_txt, re.I):
                    out.append(add_image_label_paragraph(line_txt, "Normal"))
                    continue
                if _parece_item_lista_en_bloque(line_txt, modo_lista_r):
                    out.append(p_vineta_recuerda(line if isinstance(line, dict) else _limpiar_vineta_literal(line_txt)))
                    _append_links_xml(out, line)
                    continue
                out.append(p(line, "Recuerda-Cuerpoparrafo"))
                _append_links_xml(out, line)
                modo_lista_r = _abre_modo_lista(line_txt)

            out.append(p("", "Recuerda-01lneafin"))

        elif t == "tarea":
            out.append(p(b.get("etiqueta", "Tarea"), "Recuerda-00lneainicio"))
            enunciado_start = b.get("enunciado_linea", 0)
            for i, line in enumerate(b.get("lineas", [])):
                if i < enunciado_start:
                    continue
                if isinstance(line, dict) and line.get("tipo") == "tabla":
                    out.append(tabla_xml(line.get("filas", [])))
                    continue
                if rich_text(line).strip():
                    out.append(p(line, "EjerciciosPregunta"))
                    _append_links_xml(out, line)

            for opt in b.get("opciones", []):
                out.append(p_opcion_test(opt.get("letra", ""), opt.get("texto", "")))

            # Regla editorial: las tareas y el resto de recursos prácticos no
            # llevan solución en el papel. La solución se reserva para
            # Aplicación práctica.
            out.append(p("", "Recuerda-01lneafin"))


        elif t == "aplicacion_practica":
            out.append(p(b.get("etiqueta", "Aplicación práctica"), "Recuerda-00lneainicio"))

            for line in b.get("lineas", []):
                if isinstance(line, dict) and line.get("tipo") == "tabla":
                    out.append(tabla_xml(line.get("filas", [])))
                    continue
                if rich_text(line).strip():
                    out.append(p(line, "EjerciciosPregunta"))
                    _append_links_xml(out, line)

            for opt in b.get("opciones", []):
                out.append(p_opcion_test(opt.get("letra", ""), opt.get("texto", "")))

            # Regla editorial: solo las aplicaciones prácticas llevan solución en el papel.
            for line in b.get("solucion_lineas", []):
                line_txt = rich_text(line).strip()
                if not line_txt:
                    continue
                # Strip "Feedback:" label — show the content directly without the prefix.
                m_fb = re.match(r"^(?:Feedback|Retroalimentaci[oó]n)\s*:\s*", line_txt, re.I)
                if m_fb:
                    content = line_txt[m_fb.end():]
                    if content:
                        out.append(p(rich_obj(content, [{"text": content}]), "EjerciciosRespuestas"))
                else:
                    out.append(p(line, "EjerciciosRespuestas"))

            for line in _solucion_feedback_a_lineas(b.get("solucion", ""), b.get("feedback", "")):
                out.append(p(line, "EjerciciosRespuestas"))

            out.append(p("", "Recuerda-01lneafin"))

        elif t == "actividad_complementaria":
            # La numeración no va en el encabezado del recurso: se pone al inicio
            # del enunciado.
            out.append(p("Actividad complementaria", "Recuerda-00lneainicio"))

            numero = str(b.get("numero", "")).strip()
            if not numero:
                mnum = re.search(r"(\d+)", str(b.get("etiqueta", "")))
                numero = mnum.group(1) if mnum else ""
            first = True
            for line in b.get("lineas", []):
                if isinstance(line, dict) and line.get("tipo") == "tabla":
                    out.append(tabla_xml(line.get("filas", [])))
                    continue
                if not rich_text(line).strip():
                    continue
                line2 = _normalizar_enunciado_complementaria(line) if first else line
                if not rich_text(line2).strip():
                    first = False
                    continue
                if first and numero:
                    line2 = _prefix_rich(f"{numero}. ", line2)
                out.append(p(line2, "EjerciciosPregunta"))
                _append_links_xml(out, line2)
                first = False

            out.append(p("", "Recuerda-01lneafin"))

        elif t == "url_imagen":
            out.append(p_url_imagen(b.get("url", "")))

        elif t == "url":
            out.append(p_url_recurso(b.get("url", "")))

        elif t == "pie_imagen":
            out.append(p_pie_imagen(b.get("texto", "")))

        elif t == "desc_imagen":
            out.append(p_desc_imagen(b.get("texto", "")))

        elif t == "imagen":
            if b.get("url"):
                out.append(p_url_imagen(b["url"]))
            if b.get("pie"):
                out.append(p_pie_imagen(b["pie"]))
            if b.get("descripcion"):
                out.append(p_desc_imagen(b["descripcion"]))

        elif t == "opcion_test_suelta":
            out.append(p_opcion_test(b.get("letra", ""), b.get("texto", "")))

        elif t == "parrafo_formula":
            out.append(p_formula(b.get("texto", "")))

    return out


# =============================================================================
# Gráficos / imágenes / SmartArt desde el DOCX de referencia
# =============================================================================

def _leer_document_xml(docx_path: Path) -> str:
    with zipfile.ZipFile(str(docx_path), "r") as z:
        return z.read("word/document.xml").decode("utf-8")


def _extraer_body_xml(document_xml: str) -> str:
    m = re.search(r"<w:body>([\s\S]*?)</w:body>", document_xml)
    return m.group(1) if m else ""


def _extraer_parrafos_body(body_xml: str) -> list[str]:
    """
    IMPORTANTE:
    No usar <w:p[\\s\\S]*?</w:p> porque también casa <w:pPr>.
    """
    return re.findall(r"<w:p(?:\s[^>]*)?>[\s\S]*?</w:p>", body_xml)


def _extraer_document_attrs(docx_path: Path) -> str:
    """
    Usa los namespace reales del documento de referencia.
    Esto evita perder gráficos por faltar xmlns:wp14, wpc, wpg, wps, etc.
    """
    fallback = (
        'xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" '
        'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
        'xmlns:o="urn:schemas-microsoft-com:office:office" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
        'xmlns:v="urn:schemas-microsoft-com:vml" '
        'xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" '
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:w10="urn:schemas-microsoft-com:office:word" '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" '
        'xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml" '
        'xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" '
        'xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" '
        'xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" '
        'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" '
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture" '
        'xmlns:dgm="http://schemas.openxmlformats.org/drawingml/2006/diagram" '
        'mc:Ignorable="w14 w15 wp14"'
    )

    if not docx_path.exists() or not _es_zip(docx_path):
        return fallback

    try:
        xml = _leer_document_xml(docx_path)
    except Exception:
        return fallback

    m = re.search(r"<w:document\s+([^>]*)>", xml)
    if not m:
        return fallback

    attrs = m.group(1)

    needed = {
        "xmlns:wp14": 'xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"',
        "xmlns:a": 'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"',
        "xmlns:pic": 'xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"',
        "xmlns:dgm": 'xmlns:dgm="http://schemas.openxmlformats.org/drawingml/2006/diagram"',
    }

    for key, val in needed.items():
        if key not in attrs:
            attrs += " " + val

    if "mc:Ignorable" not in attrs:
        attrs += ' mc:Ignorable="w14 w15 wp14"'

    return attrs


def _texto_plano_parrafo_xml(par_xml: str) -> str:
    textos = re.findall(r"<w:t[^>]*>(.*?)</w:t>", par_xml)
    txt = " ".join(textos)
    txt = re.sub(r"<[^>]+>", "", txt)
    txt = (
        txt.replace("&amp;", "&")
           .replace("&lt;", "<")
           .replace("&gt;", ">")
           .replace("&quot;", '"')
    )
    txt = re.sub(r"\s+", " ", txt).strip()
    return txt


def _normalizar_contexto(txt: str) -> str:
    txt = txt or ""
    txt = txt.lower()
    txt = (
        txt.replace("á", "a")
           .replace("é", "e")
           .replace("í", "i")
           .replace("ó", "o")
           .replace("ú", "u")
           .replace("ü", "u")
           .replace("ñ", "n")
    )
    txt = re.sub(r"\s+", " ", txt)
    txt = re.sub(r"[^\w ]+", "", txt)
    return txt.strip()


def _parrafo_tiene_grafico(par_xml: str) -> bool:
    """
    Detecta cualquier gráfico insertado en Word:
      - imagen normal: <a:blip r:embed="...">
      - SmartArt / diagrama: <dgm:relIds ...>
      - formas / dibujos: <w:drawing>
      - VML antiguo: <w:pict>
    """
    return (
        "<w:drawing" in par_xml
        or "<w:pict" in par_xml
        or "<a:graphicData" in par_xml
        or "<dgm:relIds" in par_xml
        or "<a:blip" in par_xml
        or "r:embed=" in par_xml
        or "r:link=" in par_xml
    )


def extraer_graficos_con_contexto(docx_path: Path) -> list[dict]:
    """
    Extrae todos los gráficos del DOCX de referencia:
      - imágenes normales
      - SmartArt
      - diagramas
      - formas
      - dibujos

    Copia el párrafo <w:drawing> completo.
    """
    if not docx_path.exists() or not _es_zip(docx_path):
        return []

    try:
        document_xml = _leer_document_xml(docx_path)
    except Exception:
        return []

    body = _extraer_body_xml(document_xml)
    if not body:
        return []

    paras = _extraer_parrafos_body(body)

    graficos = []

    ordinal_grafico = -1

    for i, par in enumerate(paras):
        if not _parrafo_tiene_grafico(par):
            continue
        ordinal_grafico += 1

        prev_txt = ""
        next_txt = ""

        for j in range(i - 1, -1, -1):
            t = _texto_plano_parrafo_xml(paras[j])
            if t:
                prev_txt = t
                break

        for j in range(i + 1, len(paras)):
            t = _texto_plano_parrafo_xml(paras[j])
            if t:
                next_txt = t
                break

        # Los gráficos que pertenecen a pantallas interactivas del online
        # suelen ir precedidos por el rótulo "Interacción X". En el libro
        # papel esas interacciones se convierten a texto/listas, por lo que
        # NO debe conservarse la captura/esquema interactivo. Sí se conservan
        # los esquemas e imágenes editoriales reales de la unidad.
        contexto_grafico = " ".join([
            prev_txt or "",
            next_txt or "",
            _texto_plano_parrafo_xml(par) or "",
        ])
        if re.search(r"\bInteracci[oó]n\s+\d+\b", contexto_grafico, re.I):
            continue

        # Las fotografías/capturas del online van junto a una URL de banco de
        # imágenes. Para papel se conserva la URL, el pie y la descripción, no
        # la imagen embebida.
        if _es_url_imagen(prev_txt) or _es_url_imagen(next_txt):
            continue

        # También se omiten miniaturas o recursos externos enlazados, como
        # vídeos/redirectores, porque deben quedar como enlace visible.
        if RE_URL.match((prev_txt or "").strip()) or RE_URL.match((next_txt or "").strip()):
            continue

        # Un gráfico sin contexto posterior fiable suele ser residuo final de
        # pantalla/documento. No lo insertamos para evitar que caiga al final o
        # en una zona incorrecta.
        if not (next_txt or "").strip():
            continue

        # Strip source pStyle so graphics don't carry foreign styles into the output.
        par_sin_estilo = re.sub(r'<w:pStyle w:val="[^"]+"/>', '', par)
        graficos.append({
            "xml": "    " + par_sin_estilo,
            "prev": prev_txt,
            "next": next_txt,
            "prev_norm": _normalizar_contexto(prev_txt),
            "next_norm": _normalizar_contexto(next_txt),
            "insertado": False,
            "ordinal_grafico": ordinal_grafico,
        })

    _preparar_rasteres_graficos(graficos, docx_path)
    _aplicar_raster_a_graficos(graficos)
    return graficos


def _contexto_grafico_util(ctx_norm: str) -> bool:
    """
    Indica si un contexto normalizado sirve como ancla fiable para recolocar un gráfico.

    En versiones anteriores se usaban también anclas muy cortas o genéricas
    como "Ejemplo", "Nota" o "Consejo". Eso podía mover un esquema a la
    primera aparición de ese rótulo en lugar de mantenerlo cerca de su página
    original.
    """
    ctx = (ctx_norm or "").strip()
    if len(ctx) < 45:
        return False

    genericos = {
        "ejemplo", "nota", "consejo", "importante", "recuerda",
        "sabias que", "video", "para saber mas", "hilo conductor",
        "cambio de pantalla", "actividad complementaria",
        "actividad colaborativa", "actividad de evaluacion", "tarea de evaluacion",
    }
    if ctx in genericos:
        return False

    # Evita anclas técnicas o residuales del online.
    if ctx.startswith(("imagen ", "interaccion ", "instruccion ")):
        return False
    if ctx.startswith(("http ", "https ")):
        return False

    return True


def _claves_contexto_grafico(ctx_norm: str) -> list[str]:
    """Devuelve claves de búsqueda de mayor a menor longitud, solo si son útiles."""
    if not _contexto_grafico_util(ctx_norm):
        return []

    max_len = min(len(ctx_norm), 220)
    longitudes = [220, 180, 140, 100, 70, 45]
    claves = []
    for n in longitudes:
        n = min(n, max_len)
        if n >= 45:
            k = ctx_norm[:n].strip()
            if k and k not in claves:
                claves.append(k)
    return claves


def _par_coincide_con_contexto(par_norm: str, ctx_norm: str) -> bool:
    claves = _claves_contexto_grafico(ctx_norm)
    return any(k and (par_norm.startswith(k) or k in par_norm) for k in claves)


def insertar_graficos_por_contexto(pars: list[str], graficos: list[dict]) -> list[str]:
    """
    Recoloca gráficos del ejemplo dentro del documento generado.

    Prioridad corregida:
      1. Insertar después del texto anterior, si el contexto es específico.
      2. Insertar antes del texto posterior, si el contexto es específico.
      3. Insertar los no localizados al final.

    No se usan anclas genéricas como "Ejemplo", "Nota" o "Consejo", porque
    aparecen muchas veces y provocan desplazamientos incorrectos de esquemas.
    """
    if not graficos:
        return pars

    # Por si se reutiliza la misma lista de gráficos en pruebas.
    for g in graficos:
        g["insertado"] = False

    resultado = []

    for par in pars:
        par_txt = _texto_plano_parrafo_xml(par)
        par_norm = _normalizar_contexto(par_txt)

        resultado.append(par)

        # 1) Ancla preferente: texto anterior del gráfico original.
        for g in graficos:
            if g["insertado"]:
                continue

            prev_norm = g.get("prev_norm", "")
            if _par_coincide_con_contexto(par_norm, prev_norm):
                resultado.append(g["xml"])
                g["insertado"] = True

        # 2) Ancla secundaria: texto posterior del gráfico original.
        # Se comprueba después de añadir el párrafo actual. Si coincide, se mueve
        # el gráfico antes de ese párrafo extrayéndolo del final de resultado.
        pendientes_antes = []
        for g in graficos:
            if g["insertado"]:
                continue

            next_norm = g.get("next_norm", "")
            if _par_coincide_con_contexto(par_norm, next_norm):
                pendientes_antes.append(g)

        if pendientes_antes:
            # Quita el párrafo actual, inserta gráficos, y vuelve a ponerlo.
            actual = resultado.pop()
            for g in pendientes_antes:
                resultado.append(g["xml"])
                g["insertado"] = True
            resultado.append(actual)

    # Los gráficos sin ancla fiable NO se fuerzan al final. Forzarlos fue lo que
    # provocó que elementos residuales o esquemas sin contexto aparecieran en una
    # zona incorrecta. Se informa por consola, pero se omiten.
    no_insertados = [g for g in graficos if not g.get("insertado")]
    if no_insertados:
        print(f"  Aviso: {len(no_insertados)} gráfico(s)/esquema(s) omitidos por no tener ancla fiable")

    return resultado




def _ids_relacion_referenciados(xml: str) -> set[str]:
    """Devuelve rIds usados por dibujos/imágenes/SmartArt insertados en document.xml."""
    ids: set[str] = set()
    # Relaciones típicas en DrawingML y SmartArt: r:embed, r:link, r:dm, r:lo, r:qs, r:cs, etc.
    for m in re.finditer(r'\br:[A-Za-z0-9_]+="([^"]+)"', xml or ""):
        ids.add(m.group(1))
    return ids


def _parse_relationships_xml(xml: bytes | str) -> list[dict]:
    txt = xml.decode('utf-8') if isinstance(xml, bytes) else str(xml or '')
    rels: list[dict] = []
    for m in re.finditer(r'<Relationship\s+([^>]*)/>', txt):
        attrs = dict(re.findall(r'([A-Za-z_:][\w:.-]*)="([^"]*)"', m.group(1)))
        if attrs.get('Id'):
            rels.append(attrs)
    return rels


def _rels_to_xml(rels: list[dict]) -> bytes:
    partes = ['<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
              '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">']
    for r in rels:
        attrs = []
        for k in ('Id', 'Type', 'Target', 'TargetMode'):
            if r.get(k):
                attrs.append(f'{k}="{esc(r[k])}"')
        partes.append('<Relationship ' + ' '.join(attrs) + '/>')
    partes.append('</Relationships>')
    return ''.join(partes).encode('utf-8')


def _nuevo_rid_libre(existing: set[str], prefix: str = 'rIdSrc') -> str:
    i = 1
    while f'{prefix}{i}' in existing:
        i += 1
    rid = f'{prefix}{i}'
    existing.add(rid)
    return rid


def _merge_content_types(archivos: dict[str, bytes], source_docx: Path) -> None:
    """Añade al content types de salida los Default/Override necesarios de la unidad fuente."""
    if not source_docx or not source_docx.exists() or not _es_zip(source_docx):
        return
    try:
        with zipfile.ZipFile(str(source_docx), 'r') as zs:
            src_ct = zs.read('[Content_Types].xml').decode('utf-8')
    except Exception:
        return

    dst_ct = archivos.get('[Content_Types].xml', b'').decode('utf-8', errors='ignore')
    if not dst_ct or '</Types>' not in dst_ct:
        return

    inserts: list[str] = []
    existing_exts = set(re.findall(r'<Default\s+[^>]*Extension="([^"]+)"', dst_ct))
    for m in re.finditer(r'<Default\s+[^>]*Extension="([^"]+)"[^>]*/>', src_ct):
        ext = m.group(1)
        if ext not in existing_exts:
            inserts.append(m.group(0))
            existing_exts.add(ext)

    existing_parts = set(re.findall(r'<Override\s+[^>]*PartName="([^"]+)"', dst_ct))
    for m in re.finditer(r'<Override\s+[^>]*PartName="([^"]+)"[^>]*/>', src_ct):
        part = m.group(1)
        # Solo interesan recursos gráficos/diagramas; no sobreescribimos document.xml ni estilos.
        if not (part.startswith('/word/diagrams/') or part.startswith('/word/media/') or part.startswith('/word/charts/') or part.startswith('/word/embeddings/')):
            continue
        if part not in existing_parts:
            inserts.append(m.group(0))
            existing_parts.add(part)

    if inserts:
        dst_ct = dst_ct.replace('</Types>', ''.join(inserts) + '</Types>')
        archivos['[Content_Types].xml'] = dst_ct.encode('utf-8')


def _copiar_recursos_graficos_fuente(archivos: dict[str, bytes], source_docx: Path) -> None:
    """Copia imágenes, diagramas, gráficos y embeddings de la unidad fuente al paquete de salida."""
    if not source_docx or not source_docx.exists() or not _es_zip(source_docx):
        return
    prefixes = ('word/media/', 'word/diagrams/', 'word/charts/', 'word/embeddings/')
    try:
        with zipfile.ZipFile(str(source_docx), 'r') as zs:
            for name in zs.namelist():
                if name.startswith(prefixes):
                    # Puede sobreescribir diagramas del ejemplo: no los usamos como contenido en v15.
                    archivos[name] = zs.read(name)
    except Exception:
        return
    _merge_content_types(archivos, source_docx)


def _fusionar_rels_graficos(document_xml: str, archivos: dict[str, bytes], source_docx: Path) -> str:
    """
    Añade a document.xml.rels las relaciones de la unidad fuente necesarias para
    los dibujos/esquemas insertados. Si un rId ya existe con otro destino, se crea
    un rId nuevo y se reemplaza en document.xml.
    """
    if not source_docx or not source_docx.exists() or not _es_zip(source_docx):
        return document_xml

    needed = _ids_relacion_referenciados(document_xml)
    if not needed:
        return document_xml

    try:
        with zipfile.ZipFile(str(source_docx), 'r') as zs:
            src_rels_xml = zs.read('word/_rels/document.xml.rels')
    except Exception:
        return document_xml

    dst_key = 'word/_rels/document.xml.rels'
    src_rels = {r['Id']: r for r in _parse_relationships_xml(src_rels_xml)}
    dst_rels_list = _parse_relationships_xml(archivos.get(dst_key, b''))
    dst_by_id = {r['Id']: r for r in dst_rels_list}
    existing = set(dst_by_id)

    replacements: dict[str, str] = {}

    for rid in sorted(needed):
        if rid not in src_rels:
            continue
        src = dict(src_rels[rid])
        if rid in dst_by_id:
            dst = dst_by_id[rid]
            same = (dst.get('Type') == src.get('Type') and dst.get('Target') == src.get('Target') and dst.get('TargetMode', '') == src.get('TargetMode', ''))
            if same:
                continue
            new_id = _nuevo_rid_libre(existing)
            replacements[rid] = new_id
            src['Id'] = new_id
            dst_rels_list.append(src)
            dst_by_id[new_id] = src
        else:
            existing.add(rid)
            dst_rels_list.append(src)
            dst_by_id[rid] = src

    # Sustituye rIds en cualquier atributo r:*="old".
    for old, new in replacements.items():
        document_xml = re.sub(rf'(\br:[A-Za-z0-9_]+=")({re.escape(old)})(")', rf'\g<1>{new}\g<3>', document_xml)

    archivos[dst_key] = _rels_to_xml(dst_rels_list)
    _copiar_recursos_graficos_fuente(archivos, source_docx)
    return document_xml

# =============================================================================
# Generación DOCX
# =============================================================================

def extraer_sectpr(docx_path: Path) -> str:
    if not docx_path.exists() or not _es_zip(docx_path):
        return "<w:sectPr/>"

    try:
        xml = _leer_document_xml(docx_path)
    except Exception:
        return "<w:sectPr/>"

    matches = re.findall(r"<w:sectPr[\s\S]*?</w:sectPr>", xml)
    if matches:
        return matches[-1]

    m = re.search(r"<w:sectPr[^>]*/>", xml)
    if m:
        return m.group(0)

    return "<w:sectPr/>"


def crear_docx_minimo() -> dict[str, bytes]:
    content_types = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        '<Override PartName="/word/styles.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>'
        "</Types>"
    ).encode("utf-8")

    rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
        'Target="word/document.xml"/>'
        "</Relationships>"
    ).encode("utf-8")

    word_rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" '
        'Target="styles.xml"/>'
        "</Relationships>"
    ).encode("utf-8")

    styles = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:style w:type="paragraph" w:default="1" w:styleId="Normal">'
        '<w:name w:val="Normal"/>'
        '<w:rPr><w:sz w:val="24"/></w:rPr>'
        "</w:style>"
        "</w:styles>"
    ).encode("utf-8")

    return {
        "[Content_Types].xml": content_types,
        "_rels/.rels": rels,
        "word/_rels/document.xml.rels": word_rels,
        "word/styles.xml": styles,
    }


def _cargar_paquete_base(base: Path | None, ejemplo: Path) -> dict[str, bytes]:
    """
    Carga el paquete base con la prioridad correcta de estilos.

    Orden de carga:
      1. ejemplo  → aporta media, relaciones y su styles.xml como base
      2. plantilla (base) → si es distinta del ejemplo, sus archivos ganan
         (incluido styles.xml), salvo word/document.xml que siempre se genera.
    Así un ejemplo maquetado distinto aplica sus propios estilos.
    """
    archivos: dict[str, bytes] = {}

    # Paso 1: cargar el ejemplo como punto de partida (media, rels, estilos del ejemplo)
    if ejemplo and ejemplo.exists() and _es_zip(ejemplo):
        with zipfile.ZipFile(str(ejemplo), "r") as zej:
            for name in zej.namelist():
                if name == "word/document.xml":
                    continue
                archivos[name] = zej.read(name)
    else:
        archivos = crear_docx_minimo()

    # Paso 2: si hay una plantilla separada, sus archivos tienen prioridad sobre
    # los del ejemplo (así los estilos de la plantilla maquetada siempre ganan)
    if base and base.exists() and _es_zip(base) and base != ejemplo:
        with zipfile.ZipFile(str(base), "r") as zin:
            for name in zin.namelist():
                if name == "word/document.xml":
                    continue
                archivos[name] = zin.read(name)

    return archivos


def _ensure_minimal_styles(archivos: dict[str, bytes]) -> None:
    """Crea estilos básicos si la plantilla no los trae. No sustituye estilos existentes."""
    key = "word/styles.xml"
    xml = archivos.get(key, b"").decode("utf-8", errors="ignore")
    if not xml or "</w:styles>" not in xml:
        archivos[key] = crear_docx_minimo()[key]
        xml = archivos[key].decode("utf-8", errors="ignore")

    specs = {
        "TITULOUNIDAD1": ("_TITULO UNIDAD 1", "32", True),
        "TITULOUNIDAD2": ("_TITULO UNIDAD 2", "28", True),
        "1Titulonvl1": ("1 Título nvl1", "26", True),
        "2Titulonvl2": ("2 Título nvl2", "24", True),
        "3Titulonvl3": ("3 Título nvl3", "23", True),
        "Cuerpoparrafo": ("Cuerpo parrafo", "22", False),
        "Vietanvl11d": ("Viñeta nvl1 1d", "22", False),
        "Vietanvl21d": ("Viñeta nvl2 1d", "22", False),
        "Ejemplos-Cuerpoparrafo": ("Ejemplos - Cuerpo párrafo", "22", False),
        "Ejemplos-Vietanvl1": ("Ejemplos - Viñeta nvl1", "22", False),
        "Ejemplos-Vietanvl2": ("Ejemplos - Viñeta nvl2", "22", False),
        "Ejemplos-01lneainicio": ("Ejemplos - 01 línea inicio", "22", True),
        "Ejemplos-02lneafin": ("Ejemplos - 02 línea fin", "8", False),
        "EjerciciosPregunta": ("Ejercicios Pregunta", "22", False),
        "EjerciciosRespuestas": ("Ejercicios Respuestas", "22", False),
        "Formula": ("Formula", "22", False),
        "Recuerda-00lneainicio": ("Recuerda - 00línea inicio", "22", True),
        "Recuerda-Cuerpoparrafo": ("Recuerda - cuerpo párrafo", "22", False),
        "Recuerda-Vietanvl1": ("Recuerda - Viñeta nvl1", "22", False),
        "Recuerda-Vietanvl2": ("Recuerda - Viñeta nvl2", "22", False),
        "Recuerda-01lneafin": ("Recuerda - 01línea fin", "8", False),
    }
    inserts = []
    for style_id, (name, size, bold) in specs.items():
        if f'w:styleId="{style_id}"' in xml:
            continue
        rpr = f'<w:rPr>{"<w:b/><w:bCs/>" if bold else ""}<w:sz w:val="{size}"/><w:szCs w:val="{size}"/></w:rPr>'
        inserts.append(
            f'<w:style w:type="paragraph" w:styleId="{style_id}">'
            f'<w:name w:val="{esc(name)}"/><w:basedOn w:val="Normal"/>'
            f'<w:qFormat/>{rpr}</w:style>'
        )
    if inserts:
        xml = xml.replace("</w:styles>", "".join(inserts) + "</w:styles>")
        archivos[key] = xml.encode("utf-8")


def generar_docx(est: dict, ejemplo: Path, plantilla: Path, salida: Path, unidad: Path | None = None, config: dict | None = None):
    if config is not None:
        set_runtime_config(config)
    _reset_hyperlinks_out()
    ns = _extraer_document_attrs(ejemplo)

    pars: list[str] = []

    # Normalizar siempre a "Unidad de aprendizaje N" para TITULOUNIDAD1
    titulo_raw = est.get("titulo_unidad", "Unidad de aprendizaje 1")
    m_num = re.search(r"\d+", titulo_raw)
    titulo_unidad_papel = f"Unidad de aprendizaje {m_num.group(0)}" if m_num else titulo_raw
    pars.append(p(titulo_unidad_papel, "TITULOUNIDAD1"))
    pars.append(p(est.get("titulo_modulo", ""), "TITULOUNIDAD2", negrita=True))

    objetivo_general = est.get("objetivo_general", "")
    objetivos = est.get("objetivos", [])

    if objetivo_general:
        pars.append(p("El objetivo general de esta Unidad de Aprendizaje es:", "Cuerpoparrafo"))
        pars.append(p(
            objetivo_general if isinstance(objetivo_general, dict) else rich_obj(objetivo_general, [{"text": objetivo_general}]),
            "Cuerpoparrafo",
        ))

    if objetivos:
        obj_intro = est.get("_objectives_intro") or RUNTIME_CONFIG.get("objectives_intro", "Los objetivos específicos de esta Unidad de Aprendizaje son:")
        pars.append(p(obj_intro, "Cuerpoparrafo"))
        for obj in objetivos:
            pars.append(p_vineta(obj, 1))

    # 1. Procesar y agrupar recursos en todos los niveles
    for sec in est.get("secciones", []):
        sec["bloques"] = remove_empty_example_blocks(sec.get("bloques", []))

        for sub in sec.get("subsecciones", []):
            sub["bloques"] = remove_empty_example_blocks(sub.get("bloques", []))

            for sub2 in sub.get("subsecciones", []):
                sub2["bloques"] = remove_empty_example_blocks(sub2.get("bloques", []))

    # 2. Generar el XML de la estructura limpia
    for sec in est.get("secciones", []):
        pars.append(p(f'{sec.get("num", "")}. {sec.get("titulo", "")}', "1Titulonvl1"))
        pars.extend(bloques_xml(sec.get("bloques", [])))

        for sub in sec.get("subsecciones", []):
            sub_titulo = sub.get("titulo", "")
            sub_num = sub.get("num", "")
            if RE_SEC2.match(sub_titulo) or RE_SEC1.match(sub_titulo):
                pars.append(p(sub_titulo, "2Titulonvl2"))
            else:
                pars.append(p(f'{sub_num} {sub_titulo}'.strip(), "2Titulonvl2"))
            pars.extend(bloques_xml(sub.get("bloques", [])))

            for sub2 in sub.get("subsecciones", []):
                sub2_titulo = sub2.get("titulo", "")
                sub2_num = sub2.get("num", "")
                if RE_SEC2.match(sub2_titulo) or RE_SEC1.match(sub2_titulo):
                    pars.append(p(sub2_titulo, "2Titulonvl2"))
                else:
                    pars.append(p(f'{sub2_num} {sub2_titulo}'.strip(), "2Titulonvl2"))
                pars.extend(bloques_xml(sub2.get("bloques", [])))


    # v15: los gráficos/esquemas se toman de la UNIDAD fuente, no del ejemplo.
    # El ejemplo sigue siendo plantilla de estilos. Así evitamos contaminar la
    # salida con esquemas de otra unidad y, a la vez, mantenemos los esquemas
    # reales del online en su posición por contexto.
    fuente_graficos = unidad if unidad and unidad.exists() and _es_zip(unidad) else None
    graficos = extraer_graficos_con_contexto(fuente_graficos) if fuente_graficos else []
    if graficos:
        pars = insertar_graficos_por_contexto(pars, graficos)
        print(f"  Gráficos/esquemas de la unidad insertados: {sum(1 for g in graficos if g.get('insertado'))}/{len(graficos)}")
    else:
        print("  Gráficos/esquemas de la unidad: 0")

    sectpr = extraer_sectpr(ejemplo)

    document_xml_str = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        f"<w:document {ns}><w:body>\n"
        + "\n".join(pars)
        + "\n"
        + sectpr
        + "\n</w:body></w:document>"
    )

    base = plantilla if plantilla and plantilla.exists() and _es_zip(plantilla) else ejemplo
    archivos = _cargar_paquete_base(base, ejemplo)
    _ensure_minimal_styles(archivos)
    if fuente_graficos and graficos:
        document_xml_str = _fusionar_rels_graficos(document_xml_str, archivos, fuente_graficos)
        _registrar_rasteres_en_paquete(archivos, graficos)
    document_xml = document_xml_str.encode("utf-8")
    archivos["word/document.xml"] = document_xml
    if "word/_rels/document.xml.rels" in archivos:
        archivos["word/_rels/document.xml.rels"] = _patch_document_rels_hyperlinks(archivos["word/_rels/document.xml.rels"])

    salida.parent.mkdir(parents=True, exist_ok=True)

    with zipfile.ZipFile(str(salida), "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in archivos.items():
            zout.writestr(name, data)


# =============================================================================
# Validación de salida
# =============================================================================

def _docx_texto_plano(path: Path) -> str:
    if not path.exists() or not _es_zip(path):
        return ""
    try:
        xml = _leer_document_xml(path)
    except Exception:
        return ""
    xml = re.sub(r"<w:tab[^>]*/>", "\t", xml)
    texts = re.findall(r"<w:t(?:\s[^>]*)?>(.*?)</w:t>", xml, flags=re.S)
    import html
    return "\n".join(html.unescape(re.sub(r"<[^>]+>", "", t)) for t in texts)


def validate_output(path: Path, cfg: dict | None = None) -> dict:
    cfg = cfg or RUNTIME_CONFIG
    text = _docx_texto_plano(path)
    errores: list[str] = []
    warnings: list[str] = []

    common = sorted(set(m.group(0) for m in RE_FORBIDDEN_COMMON.finditer(text)))
    if common:
        errores.append("Marcadores digitales sin limpiar: " + ", ".join(common[:10]))

    if cfg.get("version") == "alumno":
        student = sorted(set(m.group(0) for m in RE_FORBIDDEN_STUDENT.finditer(text)))
        if student:
            errores.append("Marcadores docentes visibles en versión alumno: " + ", ".join(student[:10]))

    nums = []
    for line in text.splitlines():
        m = RE_SEC1.match(line.strip())
        if m:
            try:
                nums.append(int(m.group(1)))
            except Exception:
                pass
    if nums:
        expected = list(range(nums[0], nums[0] + len(nums)))
        if nums != expected:
            warnings.append(f"Numeración de apartados posiblemente no correlativa: {nums[:20]}")

    if re.search(r"(?im)^Ejemplo\s*\n\s*Ejemplo", text):
        warnings.append("Se han detectado posibles bloques 'Ejemplo' consecutivos o vacíos.")

    if "http://" not in text and "https://" not in text:
        warnings.append("No se han detectado URLs en la salida; revisa si el documento fuente tenía recursos externos.")

    return {
        "errores": errores,
        "warnings": warnings,
        "resumen": {
            "version": cfg.get("version"),
            "apartados_detectados": len(nums),
            "urls_detectadas": len(re.findall(r"https?://", text)),
        },
    }


def imprimir_validacion(resultado: dict) -> None:
    print("→ Validación:")
    for err in resultado.get("errores", []):
        print(f"  ERROR: {err}")
    for warn in resultado.get("warnings", []):
        print(f"  Aviso: {warn}")
    if not resultado.get("errores") and not resultado.get("warnings"):
        print("  Sin incidencias detectadas")
    resumen = resultado.get("resumen", {})
    print(f"  Resumen: versión={resumen.get('version')}, apartados={resumen.get('apartados_detectados')}, urls={resumen.get('urls_detectadas')}")


# =============================================================================
# Main
# =============================================================================

def buscar_interacciones_auto(unidad: Path) -> Path | None:
    candidatos = sorted(unidad.parent.glob("interacciones_*.docx"))
    return candidatos[0] if candidatos else None


def _eliminar_primera_imagen_shutterstock(est: dict) -> None:
    if not est.get("secciones"):
        return
    
    primera_sec = est["secciones"][0]
    # Comprobar que es la sección de portada/Introducción (antes del apartado 1)
    if primera_sec.get("num") not in {"", None}:
        return
        
    bloques = primera_sec.get("bloques", [])
    
    idx_borrar = -1
    for i, b in enumerate(bloques):
        url = str(b.get("url", "")) or rich_text(b.get("texto"))
        if b.get("tipo") in {"url_imagen", "url"} and "shutterstock.com" in url.lower():
            idx_borrar = i
            break
            
    if idx_borrar >= 0:
        bloques.pop(idx_borrar)
        while idx_borrar < len(bloques):
            if bloques[idx_borrar].get("tipo") in {"pie_imagen", "desc_imagen"}:
                bloques.pop(idx_borrar)
            else:
                break


def main():
    raw_args = sys.argv[1:]
    version = None
    config_path = None
    validar = True
    args = []

    i = 0
    while i < len(raw_args):
        a = raw_args[i]
        if a == "--version" and i + 1 < len(raw_args):
            version = raw_args[i + 1].strip().lower()
            i += 2
        elif a.startswith("--version="):
            version = a.split("=", 1)[1].strip().lower()
            i += 1
        elif a == "--config" and i + 1 < len(raw_args):
            config_path = Path(raw_args[i + 1])
            i += 2
        elif a.startswith("--config="):
            config_path = Path(a.split("=", 1)[1])
            i += 1
        elif a == "--no-validate":
            validar = False
            i += 1
        else:
            args.append(a)
            i += 1

    if version and version not in {"alumno", "docente"}:
        print("ERROR: --version debe ser 'alumno' o 'docente'")
        sys.exit(1)

    cfg = cargar_config(config_path, version)
    set_runtime_config(cfg)

    inter_path = None

    if len(args) == 2:
        unidad, ejemplo = (Path(x) for x in args)
        plantilla = ejemplo
        salida = unidad.with_name(unidad.stem + "_resultado.docx")

    elif len(args) == 3:
        unidad, ejemplo, salida = (Path(x) for x in args)
        plantilla = ejemplo

    elif len(args) == 4:
        unidad = Path(args[0])
        ejemplo = Path(args[1])
        inter_path = Path(args[2])
        salida = Path(args[3])
        plantilla = ejemplo

    elif len(args) == 5:
        unidad = Path(args[0])
        ejemplo = Path(args[1])
        plantilla = Path(args[2])
        inter_path = Path(args[3])
        salida = Path(args[4])

    else:
        print("Uso:")
        print("  python conversor_papel_generico.py [--version alumno|docente] [--config config.json] UNIDAD.docx EJEMPLO.docx SALIDA.docx")
        print("  python conversor_papel_generico.py [--version alumno|docente] UNIDAD.docx EJEMPLO.docx INTERACCIONES.docx SALIDA.docx")
        print("  python conversor_papel_generico.py [--version alumno|docente] UNIDAD.docx EJEMPLO.docx PLANTILLA.docx INTERACCIONES.docx SALIDA.docx")
        print("  python conversor_papel_generico.py --no-validate ...")
        sys.exit(1)

    for path, label in [(unidad, "Unidad"), (ejemplo, "Ejemplo maquetado")]:
        if not path.exists():
            print(f"ERROR: no existe {label}: {path}")
            sys.exit(1)

    if inter_path is None and unidad.suffix.lower() in {".docx", ".doc"}:
        inter_path = buscar_interacciones_auto(unidad)
        if inter_path:
            print(f"→ Interacciones detectadas: {inter_path.name}")

    interacciones = {}

    if inter_path and inter_path.exists():
        print(f"→ Parseando interacciones: {inter_path.name}")
        interacciones = parsear_interacciones(inter_path)
        print(f"  {len(interacciones)} interacciones cargadas")

    print(f"→ Configuración: versión={cfg.get('version')}, soluciones={cfg.get('include_solutions')}, feedback={cfg.get('include_feedback')}")
    print(f"→ Parseando unidad: {unidad.name}")

    if unidad.suffix.lower() in {".docx", ".doc"} and _es_zip(unidad):
        est = parsear_docx_fuente(unidad, interacciones)
    else:
        est = parsear_pdf_o_texto(unidad, interacciones)

    _eliminar_primera_imagen_shutterstock(est)

    print(f'  {est.get("titulo_unidad", "")} — {est.get("titulo_modulo", "")}')
    print(f'  Objetivos: {len(est.get("objetivos", []))}')
    print(f'  Secciones: {len(est.get("secciones", []))}')

    total_bloques = 0

    for sec in est.get("secciones", []):
        nb = len(sec.get("bloques", []))
        ns = len(sec.get("subsecciones", []))
        nsb = sum(len(s.get("bloques", [])) for s in sec.get("subsecciones", []))
        total_bloques += nb + nsb
        print(f'    {sec.get("num")}. {sec.get("titulo")} [{nb} bloques, {ns} subsecs, {nsb} bloques subsec]')

    print(f"  Total bloques: {total_bloques}")

    print(f"→ Generando salida: {salida.name}")
    generar_docx(est, ejemplo, plantilla, salida, unidad, cfg)
    print(f"✓ Listo: {salida}")

    if validar:
        resultado = validate_output(salida, cfg)
        imprimir_validacion(resultado)
        if resultado.get("errores") and cfg.get("validation", {}).get("fail_on_errors"):
            sys.exit(2)


if __name__ == "__main__":
    main()
